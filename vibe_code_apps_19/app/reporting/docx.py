"""Render Engineering Findings DOCX via python-docx (optional dependency)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from app.reporting.models import Classification, ReportArtifacts
from app.reporting.narrative import confidence_badge
from app.reporting.rule_meta import legend_rows, rule_label

def render_docx(artifacts: ReportArtifacts, path: Path | str) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for Engineering Findings DOCX. "
            "Install: pip install 'vibe19-fdd-demo[engineering-report]' or python-docx"
        ) from exc

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Cover
    title = doc.add_heading("FDD Engineering Findings Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(artifacts.building)
    run.bold = True
    run.font.size = Pt(16)
    _muted(doc, f"Analysis period: {artifacts.analysis_period or 'see assumptions'}")
    _muted(doc, f"Generated: {artifacts.generated_at}")
    disc = doc.add_paragraph()
    run = disc.add_run(
        "Open-FDD advisory analysis. Findings are telemetry-based; physical verification "
        "remains a field activity. Detection ≠ finding."
    )
    run.italic = True
    run.font.size = Pt(10)
    doc.add_page_break()

    # Executive summary
    doc.add_heading("1. Executive summary", level=1)
    m = artifacts.metrics or {}
    doc.add_paragraph(
        f"Detections reviewed: {m.get('n_candidates', '—')}. "
        f"Strongly supported: {m.get('n_strongly_supported', 0)}; "
        f"probable: {m.get('n_probable', 0)}; "
        f"inconclusive: {m.get('n_inconclusive', 0)}; "
        f"likely false positives (appendix): {m.get('n_suppressed', 0)}; "
        f"data-quality issues: {m.get('n_data_quality', 0)}."
    )
    doc.add_paragraph(
        "Raw rule hits were evidence-reviewed before appearing as prioritized findings below."
    )

    included = [f for f in artifacts.findings if f.include_in_report]
    if included:
        doc.add_heading("Priority index", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Priority"
        hdr[1].text = "Finding"
        hdr[2].text = "Status"
        hdr[3].text = "First field check"
        for f in included:
            row = table.add_row().cells
            row[0].text = str(f.priority)
            row[1].text = f.title
            row[2].text = confidence_badge(f.effective_classification)
            row[3].text = (f.field_verification or ["—"])[0]

    # Building at a glance
    doc.add_heading("2. Building at a glance", level=1)
    doc.add_paragraph(
        "Charts below are generated from this evidence review (Plotly → PNG). "
        "They mirror Overview-style summaries where data is available; interactive "
        "Overview session figures are not auto-copied into the Word file."
    )
    _add_chart_if_any(doc, artifacts, "confidence_summary")
    _add_chart_if_any(doc, artifacts, "top_detections")
    _add_chart_if_any(doc, artifacts, "comfort_ranking")

    # Prioritized findings detail
    doc.add_heading("3. Prioritized engineering findings", level=1)
    for f in included:
        doc.add_heading(f"{f.finding_id}: {f.title}", level=2)
        doc.add_paragraph(f"Status: {confidence_badge(f.effective_classification)}")
        if f.rule_ids:
            bits = [f"{rule_label(rid)} ({rid})" for rid in f.rule_ids]
            doc.add_paragraph(f"Rule(s): {', '.join(bits)}")
        doc.add_paragraph(f"Why it matters: {f.why_it_matters}")
        doc.add_paragraph(f"Observed behavior: {f.observed_behavior}")
        doc.add_paragraph("Evidence:")
        for b in f.evidence_bullets:
            doc.add_paragraph(b, style="List Bullet")
        if f.chart_path and Path(f.chart_path).is_file():
            try:
                doc.add_picture(f.chart_path, width=Inches(5.8))
            except Exception:
                pass
        doc.add_paragraph("Contradicting evidence:")
        for b in f.contradicting_evidence:
            doc.add_paragraph(b, style="List Bullet")
        doc.add_paragraph("Likely causes:")
        for b in f.likely_causes:
            doc.add_paragraph(b, style="List Bullet")
        doc.add_paragraph("Recommended field verification:")
        for b in f.field_verification:
            doc.add_paragraph(b, style="List Bullet")
        doc.add_paragraph("Possible corrective action:")
        for b in f.possible_corrective:
            doc.add_paragraph(b, style="List Bullet")
        doc.add_paragraph(
            f"Equipment: {', '.join(f.equipment_ids)} · Systems: {', '.join(f.systems) or '—'}"
        )
        if f.engineer_override:
            doc.add_paragraph(
                f"Engineer override: {f.engineer_override} (automated assessment retained in JSON)"
            )

    # Comfort
    doc.add_heading("4. Comfort / zone performance", level=1)
    cs = artifacts.comfort_summary or {}
    doc.add_paragraph(
        f"VAVs: {cs.get('n_vav', '—')}; below threshold: {cs.get('n_below', '—')}. "
        "Dead/implausible sensors are excluded from comfort conclusions."
    )
    _add_chart_if_any(doc, artifacts, "comfort_ranking")
    dq_zones = [
        r
        for r in artifacts.data_quality
        if "zone" in " ".join(r.get("reasons") or []).lower() or r.get("mean_zone_t")
    ]
    if dq_zones:
        doc.add_paragraph("Instrumentation exclusions (not comfort complaints):")
        for r in dq_zones[:8]:
            doc.add_paragraph(
                f"{r.get('equipment_id')}: {'; '.join(r.get('reasons') or [])}",
                style="List Bullet",
            )

    # Systems
    doc.add_heading("5. AHU / plant / system findings", level=1)
    sys_findings = [
        f for f in included if any(s in f.systems for s in ("AHU", "CHW", "HW"))
    ]
    if not sys_findings:
        doc.add_paragraph("No separate system-level findings beyond the prioritized list.")
    else:
        doc.add_paragraph("Index of prioritized findings that touch AHU / CHW / HW systems:")
        for f in sys_findings:
            doc.add_paragraph(
                f"{f.finding_id} ({', '.join(f.systems)}): {f.title}",
                style="List Bullet",
            )

    # Data quality
    doc.add_heading("6. Data quality", level=1)
    if not artifacts.data_quality:
        doc.add_paragraph("No major data-quality exclusions beyond suppressed rows.")
    for r in artifacts.data_quality[:15]:
        rid = r.get("rule_id") or ""
        label = rule_label(str(rid)) if rid else ""
        doc.add_paragraph(
            f"{r.get('equipment_id')} / {label} ({rid}): {'; '.join(r.get('reasons') or [])}",
            style="List Bullet",
        )

    # Field checklist
    doc.add_heading("7. Recommended field checklist", level=1)
    for item in artifacts.field_checklist[:12]:
        doc.add_paragraph(f"[ ] {item}", style="List Bullet")

    # Appendix A — detections with hours (not FP dump)
    doc.add_page_break()
    doc.add_heading("Appendix A — Full rule / detection results", level=1)
    doc.add_paragraph(
        "Raw FAULT detections with hours and evidence-review class. "
        "Clients should not need this appendix to understand the building. "
        "See Appendix B for rule ID → description mapping."
    )
    class_by_key = _class_by_candidate_key(artifacts)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Equipment"
    hdr[1].text = "Rule ID"
    hdr[2].text = "Description"
    hdr[3].text = "Hours"
    hdr[4].text = "Pct"
    hdr[5].text = "Class"
    ranked = sorted(
        artifacts.candidates or [],
        key=lambda c: -(float(c.get("fault_hours") or 0)),
    )
    for c in ranked[:100]:
        rid = str(c.get("rule_id") or "")
        key = f"{c.get('equipment_id')}|{rid}"
        row = table.add_row().cells
        row[0].text = str(c.get("equipment_id") or "")
        row[1].text = rid
        row[2].text = str(c.get("rule_label") or rule_label(rid))
        fh = c.get("fault_hours")
        fp = c.get("fault_pct")
        row[3].text = f"{fh:.2f}" if isinstance(fh, (int, float)) else str(fh or "")
        row[4].text = f"{fp:.2f}" if isinstance(fp, (int, float)) else str(fp or "")
        row[5].text = class_by_key.get(key, "")

    # Appendix B — rule legend
    doc.add_heading("Appendix B — Fault / rule description legend", level=1)
    doc.add_paragraph(
        "Maps Open-FDD rule IDs used in this report to human-readable titles and summaries."
    )
    rule_ids = {str(c.get("rule_id") or "") for c in (artifacts.candidates or [])}
    for f in artifacts.findings:
        rule_ids.update(f.rule_ids or [])
    for s in artifacts.suppressed or []:
        for part in str(s.get("rule_id") or "").split(","):
            rule_ids.add(part.strip())
    legend = legend_rows(rule_ids)
    if legend:
        ltab = doc.add_table(rows=1, cols=3)
        ltab.style = "Table Grid"
        lh = ltab.rows[0].cells
        lh[0].text = "Rule ID"
        lh[1].text = "Title"
        lh[2].text = "Summary"
        for row_data in legend:
            row = ltab.add_row().cells
            row[0].text = row_data["rule_id"]
            row[1].text = row_data["title"]
            row[2].text = row_data["summary"]
    else:
        doc.add_paragraph("No rules in this report.")

    # Appendix C — likely false positives (bottom)
    doc.add_heading("Appendix C — Likely false positives / suppressed", level=1)
    doc.add_paragraph(
        "Classifications suppressed from the client body (likely false positive, "
        "not actionable, or deprioritized beyond the top findings cap)."
    )
    fp_only = [
        s
        for s in (artifacts.suppressed or [])
        if "FALSE_POSITIVE" in str(s.get("classification") or "").upper()
        or str(s.get("classification") or "").upper() == "NOT_ACTIONABLE"
    ]
    other_suppressed = [s for s in (artifacts.suppressed or []) if s not in fp_only]
    if not artifacts.suppressed:
        doc.add_paragraph("None.")
    else:
        ftab = doc.add_table(rows=1, cols=4)
        ftab.style = "Table Grid"
        fh = ftab.rows[0].cells
        fh[0].text = "Equipment"
        fh[1].text = "Rule ID"
        fh[2].text = "Description"
        fh[3].text = "Class"
        for s in fp_only + other_suppressed:
            rid = str(s.get("rule_id") or "")
            desc = ", ".join(rule_label(p.strip()) for p in rid.split(",") if p.strip()) or ""
            row = ftab.add_row().cells
            row[0].text = str(s.get("equipment_id") or "")
            row[1].text = rid
            row[2].text = desc
            row[3].text = str(s.get("classification") or "")

    # Appendix D — methods
    doc.add_heading("Appendix D — Methods / assumptions", level=1)
    for k, v in (artifacts.assumptions or {}).items():
        doc.add_paragraph(f"{k}: {v}")
    qg = artifacts.quality_gate or {}
    doc.add_paragraph(f"Quality gate OK: {qg.get('ok')}; warnings: {qg.get('warnings')}")

    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = f"{artifacts.building} — FDD Engineering Findings Report (advisory)"

    doc.save(str(path))
    return path


def _muted(doc, text: str) -> None:
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)


def _class_by_candidate_key(artifacts: ReportArtifacts) -> dict[str, str]:
    out: dict[str, str] = {}
    for a in artifacts.assessments or []:
        key = str(a.get("candidate_key") or "")
        cls = a.get("classification")
        if key and cls:
            out[key] = str(cls)
    for s in artifacts.suppressed or []:
        key = str(s.get("candidate_key") or "")
        cls = s.get("classification")
        if key and cls and "|" in key:
            out.setdefault(key, str(cls))
    return out


def _add_chart_if_any(doc, artifacts: ReportArtifacts, name: str) -> None:
    for ch in artifacts.charts or []:
        if ch.get("name") == name and ch.get("path") and Path(ch["path"]).is_file():
            try:
                from docx.shared import Inches

                doc.add_picture(ch["path"], width=Inches(5.8))
            except Exception:
                pass
            return
