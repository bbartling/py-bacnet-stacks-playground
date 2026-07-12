"""DOCX FDD / data-model reports for browser download (python-docx)."""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any

import pandas as pd

from app.data_model_tree import BuildingDataModelTree, build_data_model_tree
from app.rule_card import (
    PLACE_PLOT_HERE,
    PLACE_RCX_PLOT_HERE,
    RuleCard,
    build_rule_card,
    equipment_mapping_coverage,
)
from app.column_map_json import FAMILY_LABELS, FAMILY_ORDER
from app.rule_plot_meta import RCX_PRESETS_BY_RULE, analytics_related
from app.rules import RULES, RULES_BY_ID
from app.rules.base import RuleResult
from app.rules.cookbook_catalog import CookbookRule
from app.rules.runner import infer_equipment_kind


# Engineer fills before distributing — kept identical across all DOCX builders.
KEY_FINDINGS_PLACEHOLDER = (
    "[KEY FINDINGS — engineer summary: paste top issues, savings opportunities, "
    "and follow-ups here before distributing this report.]"
)

# Feed-relationship faults shown under parent AHU in the by-system FDD pack.
FEED_RELATIONSHIP_RULE_IDS = frozenset({"VAV-AHU-LEAVE"})


def _add_key_findings_placeholder(doc) -> None:
    _add_heading(doc, "Key findings", 1)
    _add_para(doc, KEY_FINDINGS_PLACEHOLDER)


def _docx():
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx is required for DOCX reports. Install: pip install python-docx"
        ) from exc
    return Document, Inches, Pt


def _add_heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)


def applicable_rules_for_equipment(
    equipment_id: str,
    *,
    equipment_type: str = "",
    mapped_df: pd.DataFrame | None = None,
    role_map: dict | None = None,
) -> list[CookbookRule]:
    """Canonical cookbook rules applicable to this device's equipment kind."""
    kind = infer_equipment_kind(
        equipment_id,
        equipment_type=equipment_type,
        df=mapped_df,
        role_map=role_map,
    )
    if kind == "unknown":
        return list(RULES)
    return [r for r in RULES if kind in r.equipment_kinds]


def _add_params_table(doc, card: RuleCard) -> None:
    _add_para(doc, "Sliders (tune params)", bold=True)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(
        ["Key", "Label", "Unit", "Value", "Default", "Min", "Max", "Step"]
    ):
        hdr[i].text = h
    if not card.param_rows:
        cells = table.add_row().cells
        cells[0].text = "(no tune params)"
        for i in range(1, 8):
            cells[i].text = "—"
        return
    for pr in card.param_rows:
        cells = table.add_row().cells
        cells[0].text = pr.key
        cells[1].text = pr.label
        cells[2].text = pr.unit
        cells[3].text = f"{pr.value:g}"
        cells[4].text = f"{pr.default:g}"
        cells[5].text = f"{pr.min:g}"
        cells[6].text = f"{pr.max:g}"
        cells[7].text = f"{pr.step:g}"


def _add_mapping_table(doc, card: RuleCard) -> None:
    _add_para(doc, "Points → Haystack tags", bold=True)
    if card.points_note:
        _add_para(doc, card.points_note)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(
        ["Cookbook role", "Haystack-like tag", "CSV column", "Requirement", "In history"]
    ):
        hdr[i].text = h
    if not card.mapping_rows:
        cells = table.add_row().cells
        cells[0].text = "(sensor/control sweep — applies to present sensors / outputs)"
        for i in range(1, 5):
            cells[i].text = "—"
        return
    for m in card.mapping_rows:
        cells = table.add_row().cells
        cells[0].text = m.role
        cells[1].text = m.haystack_tag
        cells[2].text = m.csv_column
        cells[3].text = m.requirement
        cells[4].text = "yes" if m.in_history else "MISSING"


def _append_rule_card_section(
    doc,
    card: RuleCard,
    *,
    plot_png: bytes | None,
    Inches,
) -> None:
    _add_heading(doc, f"{card.rule_id} — {card.title} · {card.status}", 1)
    if card.description:
        _add_para(doc, f"Summary: {card.description}")
    if card.equation:
        _add_para(doc, f"Equation: {card.equation}")
    fh = "—" if card.fault_hours is None else f"{card.fault_hours:.2f}"
    facts = list(card.catalog_facts) if card.catalog_facts else [
        ("Family", card.family),
        ("Equipment kinds", ", ".join(card.equipment_kinds) or "—"),
        ("Operational gate", card.gate_mode),
        ("Default confirm", f"{card.confirm_seconds:g}s"),
        ("Sweep", card.sweep_label),
    ]
    facts.extend(
        [
            ("Status", card.status),
            ("Fault hours", fh),
            (
                "Missing roles",
                ", ".join(card.missing_roles) if card.missing_roles else "—",
            ),
            ("Notes", card.notes or "—"),
            (
                "Required mapping",
                (
                    f"{card.required_roles_present}/{card.required_roles_total}"
                    if card.required_roles_total
                    else "n/a (sweep)"
                ),
            ),
        ]
    )
    _add_para(doc, "Rule facts", bold=True)
    _add_kv_table(doc, facts)
    _add_mapping_table(doc, card)
    _add_para(doc, "Plot series", bold=True)
    if card.plot_series:
        for bullet in card.plot_series:
            _add_para(doc, f"• {bullet}")
    else:
        _add_para(doc, "—")
    _add_params_table(doc, card)
    _add_para(doc, "Analytics / related views", bold=True)
    _add_para(doc, card.analytics_hint or "—")
    for line in card.analytics_fit:
        _add_para(doc, f"• {line}")
    _add_para(doc, "Plot", bold=True)
    if plot_png:
        doc.add_picture(io.BytesIO(plot_png), width=Inches(6.0))
    else:
        _add_para(doc, PLACE_PLOT_HERE)


def build_equipment_fdd_docx(
    *,
    building_id: str,
    equipment_id: str,
    equipment_type: str,
    results: list[RuleResult],
    role_map: dict,
    mapped_df: pd.DataFrame | None,
    plot_png_by_rule: dict[str, bytes] | None = None,
    params: dict[str, Any] | None = None,
    rules: list[CookbookRule] | None = None,
    motor_weekly: pd.DataFrame | None = None,
    cool_bins: pd.DataFrame | None = None,
) -> bytes:
    """DOCX mirror of Plots rule cards: params, mapping, PLACE PLOT HERE stubs."""
    Document, Inches, _Pt = _docx()
    doc = Document()
    applicable = rules or applicable_rules_for_equipment(
        equipment_id,
        equipment_type=equipment_type,
        mapped_df=mapped_df,
        role_map=role_map,
    )
    present, total, cov_pct = equipment_mapping_coverage(
        applicable, equipment_id, role_map, mapped_df
    )
    _add_heading(doc, f"FDD validation report — {equipment_id}", 0)
    _add_key_findings_placeholder(doc)
    _add_kv_table(
        doc,
        [
            ("Building", building_id or "—"),
            ("Equipment", equipment_id),
            ("Type", equipment_type or "—"),
            (
                "Mapping coverage",
                f"{cov_pct:.0f}% required roles present ({present}/{total})",
            ),
            (
                "Generated",
                datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
            ),
            ("Rules in report", str(len(applicable))),
        ],
    )
    _add_para(
        doc,
        "Each section mirrors a Plots validation card: equation, status, tune params, "
        "required vs mapped points, and a plot stub for paste-in figures.",
    )

    by_id = {r.rule_id: r for r in results if r.equipment_id == equipment_id}
    plot_png_by_rule = plot_png_by_rule or {}

    for rule in applicable:
        card = build_rule_card(
            equipment_id=equipment_id,
            rule=rule,
            result=by_id.get(rule.id),
            role_map=role_map,
            mapped_df=mapped_df,
            params=params,
            results=results,
        )
        _append_rule_card_section(
            doc,
            card,
            plot_png=plot_png_by_rule.get(rule.id),
            Inches=Inches,
        )

    # Optional short analytics appendix (tables only)
    if motor_weekly is not None or cool_bins is not None:
        _add_heading(doc, "Analytics appendix", 1)
        if motor_weekly is not None and not motor_weekly.empty:
            _add_para(doc, "Motor weekly (summary)", bold=True)
            cols = list(motor_weekly.columns)[:6]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            for i, c in enumerate(cols):
                table.rows[0].cells[i].text = str(c)
            for _, row in motor_weekly.head(20).iterrows():
                cells = table.add_row().cells
                for i, c in enumerate(cols):
                    cells[i].text = str(row[c])
        if cool_bins is not None and not cool_bins.empty:
            _add_para(doc, "Mechanical cooling OAT bins (summary)", bold=True)
            cols = list(cool_bins.columns)[:6]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            for i, c in enumerate(cols):
                table.rows[0].cells[i].text = str(c)
            for _, row in cool_bins.head(20).iterrows():
                cells = table.add_row().cells
                for i, c in enumerate(cols):
                    cells[i].text = str(row[c])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_building_data_model_docx(tree: BuildingDataModelTree) -> bytes:
    Document, _Inches, _Pt = _docx()
    doc = Document()
    _add_heading(doc, f"Data model — {tree.building_id or 'building'}", 0)
    _add_key_findings_placeholder(doc)
    _add_para(
        doc,
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}. "
        "Cookbook roles ↔ Haystack-like tags ↔ raw CSV columns.",
    )
    for eq in tree.equipment:
        feed_bits = []
        if getattr(eq, "fed_by", None):
            feed_bits.append(f"fedBy {eq.fed_by}")
        if getattr(eq, "feeds", None):
            feed_bits.append(f"feeds {len(eq.feeds)} VAV(s)")
        title = f"{eq.equipment_id} ({eq.equipment_type})"
        if feed_bits:
            title += " — " + "; ".join(feed_bits)
        _add_heading(doc, title, 1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(
            ["Cookbook role", "Haystack tag", "CSV column", "In history", "Required by rules"]
        ):
            hdr[i].text = h
        if not eq.bindings:
            cells = table.add_row().cells
            cells[0].text = "(no bindings)"
            for i in range(1, 5):
                cells[i].text = "—"
        for b in eq.bindings:
            cells = table.add_row().cells
            cells[0].text = b.cookbook_role
            cells[1].text = b.haystack_tag
            cells[2].text = b.csv_column or "—"
            cells[3].text = "yes" if b.present_in_history else "no"
            cells[4].text = ", ".join(b.required_by_rules[:12]) + (
                "…" if len(b.required_by_rules) > 12 else ""
            )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_analytics_docx(
    *,
    building_id: str,
    motor_weekly: pd.DataFrame | None = None,
    cool_bins: pd.DataFrame | None = None,
    rcx_coverage: pd.DataFrame | None = None,
    tree: BuildingDataModelTree | None = None,
    plant_pump_summaries: dict[str, pd.DataFrame] | None = None,
    fan_on_summaries: dict[str, pd.DataFrame] | None = None,
) -> bytes:
    Document, _Inches, _Pt = _docx()
    doc = Document()
    _add_heading(doc, f"Analytics report — {building_id or 'building'}", 0)
    _add_key_findings_placeholder(doc)
    _add_para(doc, f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    _add_para(
        doc,
        "Plant leave-temp / reset summaries use pump-on / pump-off slices when pump proof "
        "roles exist. Air-side summaries use fan-on / fan-off. Empty slices mean no proof mapped.",
    )

    def _df_section(title: str, df: pd.DataFrame | None) -> None:
        _add_heading(doc, title, 1)
        if df is None or df.empty:
            _add_para(doc, "[Placeholder] No rows for this analytics view.")
            return
        cols = list(df.columns)[:8]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = str(c)
        for _, row in df.head(40).iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(cols):
                cells[i].text = str(row[c])
        if len(df) > 40:
            _add_para(doc, f"… {len(df) - 40} more rows (see CSV export in app).")

    _df_section("Motor run hours (weekly)", motor_weekly)
    _df_section("Mechanical cooling OAT bins", cool_bins)
    _df_section("RCx preset coverage", rcx_coverage)
    if fan_on_summaries:
        for mode, df in fan_on_summaries.items():
            _df_section(f"Air-side summary stats ({mode})", df)
    if plant_pump_summaries:
        for mode, df in plant_pump_summaries.items():
            _df_section(f"Plant leave-temp summary ({mode})", df)
    if tree is not None:
        _add_heading(doc, "Data model summary", 1)
        _add_para(doc, f"{len(tree.equipment)} equipment nodes in mapping tree.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_rcx_catalog_docx(
    *,
    building_id: str,
    frames: dict[str, pd.DataFrame],
    role_map: dict,
    weather: pd.DataFrame | None = None,
    results: list[RuleResult] | None = None,
    params: dict[str, Any] | None = None,
    rcx_coverage: pd.DataFrame | None = None,
    zone_lo_f: float = 70.0,
    zone_hi_f: float = 75.0,
    occupancy_schedule: dict | None = None,
    unit_system: str = "imperial",
    motor_weekly: pd.DataFrame | None = None,
    cool_bins: pd.DataFrame | None = None,
) -> bytes:
    """Building-level RCx / catalog DOCX: RULE_PLOT_CATALOG shape with analytics filled when fit."""
    from collections import defaultdict

    from app.occupancy import OccupancySchedule
    from app.role_map import apply_role_map
    from app.rcx_plots import rcx_preset_coverage as _rcx_cov
    from app.site_model import resolve_equipment_type

    Document, _Inches, _Pt = _docx()
    doc = Document()
    results = results or []
    schedule = OccupancySchedule.from_dict(occupancy_schedule)
    if rcx_coverage is None:
        rcx_coverage = _rcx_cov(
            frames,
            role_map,
            weather=weather,
            schedule=schedule,
            comfort_low_f=zone_lo_f,
            comfort_high_f=zone_hi_f,
        )

    _add_heading(doc, f"RCx catalog report — {building_id or 'building'}", 0)
    _add_key_findings_placeholder(doc)
    _add_kv_table(
        doc,
        [
            ("Building", building_id or "—"),
            ("Unit system", unit_system),
            ("Zone comfort band °F", f"{zone_lo_f:g} – {zone_hi_f:g}"),
            ("Occupancy timezone", schedule.timezone),
            ("Equipment count", str(len(frames))),
            (
                "Generated",
                datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
            ),
        ],
    )
    _add_para(
        doc,
        "Catalog-shaped sections for all 50 cookbook rules. Analytics / RCx tables are filled "
        "when the data model supports them; otherwise an explicit not-fit reason is shown. "
        "Paste Plotly exports into PLACE RCX PLOT HERE stubs.",
    )

    if rcx_coverage is not None and not rcx_coverage.empty:
        _add_heading(doc, "RCx preset coverage", 1)
        cols = [c for c in ("preset_id", "title", "series_count", "row_count", "empty_reason") if c in rcx_coverage.columns]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = str(c)
        for _, row in rcx_coverage.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(cols):
                cells[i].text = str(row[c])

    # Representative equipment per type for mapping cards
    by_type: dict[str, str] = {}
    for eq_id, raw in frames.items():
        et = resolve_equipment_type(eq_id, df=raw, role_map=role_map)
        by_type.setdefault(et, eq_id)

    by_fam: dict[str, list] = defaultdict(list)
    for rule in RULES:
        by_fam[rule.family].append(rule)

    cov_by_id = {}
    if rcx_coverage is not None and not rcx_coverage.empty:
        cov_by_id = {
            str(r.preset_id): r for r in rcx_coverage.itertuples() if hasattr(r, "preset_id")
        }

    for fam in FAMILY_ORDER:
        rules = by_fam.get(fam) or []
        if not rules:
            continue
        _add_heading(doc, FAMILY_LABELS.get(fam, fam), 1)
        for rule in rules:
            # Pick a sample device of a matching kind when possible
            sample_eq = ""
            for kind in rule.equipment_kinds:
                # map cookbook kinds roughly to equipment types
                kind_u = kind.upper().replace("HEATPUMP", "HP").replace("ZONE", "VAV")
                if kind_u == "AHU" and "AHU" in by_type:
                    sample_eq = by_type["AHU"]
                    break
                if kind_u == "VAV" and "VAV" in by_type:
                    sample_eq = by_type["VAV"]
                    break
                if kind_u in {"CHILLER", "CHW_PLANT"} and (
                    "CHILLER" in by_type or "CHW_PLANT" in by_type
                ):
                    sample_eq = by_type.get("CHILLER") or by_type.get("CHW_PLANT") or ""
                    break
                if kind_u == "BOILER" and "BOILER" in by_type:
                    sample_eq = by_type["BOILER"]
                    break
            if not sample_eq and frames:
                sample_eq = next(iter(frames))

            mapped = None
            et = ""
            if sample_eq:
                mapped = apply_role_map(frames[sample_eq], sample_eq, role_map)
                et = resolve_equipment_type(sample_eq, df=frames[sample_eq], role_map=role_map)

            res = next(
                (
                    r
                    for r in results
                    if r.equipment_id == sample_eq and r.rule_id == rule.id
                ),
                None,
            )
            card = build_rule_card(
                equipment_id=sample_eq or "—",
                rule=rule,
                result=res,
                role_map=role_map,
                mapped_df=mapped,
                params=params,
                results=results,
                rcx_coverage=rcx_coverage,
                weather=weather,
            )
            _add_heading(doc, f"{card.rule_id} — {card.title}", 2)
            if sample_eq:
                _add_para(doc, f"Sample equipment for mapping: {sample_eq} ({et})")
            _append_rule_card_section(doc, card, plot_png=None, Inches=_Inches)

            # Fill RCx analytics for linked presets
            related = analytics_related(rule.id)
            preset_ids = related.rcx_preset_ids or RCX_PRESETS_BY_RULE.get(rule.id, ())
            for pid in preset_ids:
                stub = PLACE_RCX_PLOT_HERE.format(preset_id=pid)
                _add_para(doc, stub)
                row = cov_by_id.get(pid)
                if row is None:
                    _add_para(doc, f"RCx `{pid}`: not in coverage table")
                else:
                    n = int(getattr(row, "series_count", 0) or 0)
                    rc = int(getattr(row, "row_count", 0) or 0)
                    if rc > 0:
                        _add_para(doc, f"RCx `{pid}`: fit — {n} series / {rc} rows")
                    else:
                        reason = str(getattr(row, "empty_reason", "") or "no data")
                        _add_para(doc, f"RCx `{pid}`: not fit — {reason}")

            if rule.id == "SCHED-1":
                _add_para(
                    doc,
                    f"Overview schedule timezone `{schedule.timezone}`; "
                    f"zone band {zone_lo_f:g}–{zone_hi_f:g} °F.",
                )

    if motor_weekly is not None or cool_bins is not None:
        _add_heading(doc, "Building analytics appendix", 1)
        if motor_weekly is not None and not motor_weekly.empty:
            _add_para(doc, "Motor weekly (head)", bold=True)
            cols = list(motor_weekly.columns)[:6]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            for i, c in enumerate(cols):
                table.rows[0].cells[i].text = str(c)
            for _, row in motor_weekly.head(15).iterrows():
                cells = table.add_row().cells
                for i, c in enumerate(cols):
                    cells[i].text = str(row[c])
        if cool_bins is not None and not cool_bins.empty:
            _add_para(doc, "Mech-cooling OAT bins (head)", bold=True)
            cols = list(cool_bins.columns)[:6]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            for i, c in enumerate(cols):
                table.rows[0].cells[i].text = str(c)
            for _, row in cool_bins.head(15).iterrows():
                cells = table.add_row().cells
                for i, c in enumerate(cols):
                    cells[i].text = str(row[c])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def try_rule_plot_png(
    mapped_df: pd.DataFrame,
    result: RuleResult,
    *,
    required_roles: list[str] | None = None,
    units_map: dict[str, str] | None = None,
) -> bytes | None:
    """Best-effort Plotly→PNG for DOCX; returns None if kaleido unavailable."""
    try:
        from app.charts import rule_result_chart

        fig = rule_result_chart(
            mapped_df,
            result,
            required_roles=required_roles,
            units_map=units_map,
        )
        if fig is None:
            return None
        return fig.to_image(format="png", width=900, height=480, scale=1)
    except Exception:
        return None


def build_fdd_by_system_docx(
    *,
    building_id: str,
    frames: dict[str, pd.DataFrame],
    role_map: dict,
    results: list[RuleResult],
    equipment_ids: list[str] | None = None,
    params: dict[str, Any] | None = None,
    vav_to_ahu: dict[str, str] | None = None,
    zone_lo_f: float = 70.0,
    zone_hi_f: float = 75.0,
    occupancy_schedule: dict | None = None,
) -> bytes:
    """FDD DOCX organized by mechanical system (AHU+feeds, one VAV/zone chapter, plant)."""
    from app.occupancy import OccupancySchedule
    from app.role_map import apply_role_map
    from app.rcx_plots import zone_comfort_fail_ranking
    from app.site_model import resolve_equipment_type
    from app.topology_enrich import invert_vav_to_ahu

    Document, Inches, _Pt = _docx()
    doc = Document()
    scope = set(equipment_ids) if equipment_ids is not None else set(frames.keys())
    topo = dict(vav_to_ahu or {})
    children = invert_vav_to_ahu(topo)
    by_result: dict[tuple[str, str], RuleResult] = {
        (r.equipment_id, r.rule_id): r for r in results
    }

    _add_heading(doc, f"FDD by mechanical system — {building_id or 'building'}", 0)
    _add_key_findings_placeholder(doc)
    _add_para(
        doc,
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}. "
        "AHU sections include feed-relationship faults from child VAVs. "
        "VAV boxes are aggregated into one Zones chapter (ranking + placeholders).",
    )
    _add_para(doc, "Plot", bold=True)
    _add_para(doc, PLACE_PLOT_HERE)

    def _mapped(eq_id: str) -> pd.DataFrame:
        raw = frames[eq_id]
        mapped = apply_role_map(raw, eq_id, role_map)
        if "ahu_sat" in raw.columns and "ahu_sat" not in mapped.columns:
            mapped["ahu_sat"] = raw["ahu_sat"]
        mapped.attrs.update({k: v for k, v in raw.attrs.items() if k != "_role_map"})
        mapped.attrs["equipment_id"] = eq_id
        return mapped

    def _append_cards(eq_id: str, etype: str, rule_ids: list[str] | None = None) -> None:
        mapped = _mapped(eq_id)
        applicable = applicable_rules_for_equipment(
            eq_id, equipment_type=etype, mapped_df=mapped, role_map=role_map
        )
        if rule_ids is not None:
            want = set(rule_ids)
            applicable = [r for r in applicable if r.id in want]
        for rule in applicable:
            card = build_rule_card(
                equipment_id=eq_id,
                rule=rule,
                result=by_result.get((eq_id, rule.id)),
                role_map=role_map,
                mapped_df=mapped,
                params=params,
                results=results,
            )
            _append_rule_card_section(doc, card, plot_png=None, Inches=Inches)

    # --- AHU / RTU chapters (with child feed faults) ---
    ahu_ids = sorted(
        [
            eq
            for eq in scope
            if eq in frames
            and resolve_equipment_type(eq, df=frames[eq], role_map=role_map) in {"AHU", "RTU"}
        ]
    )
    for ahu_id in ahu_ids:
        et = resolve_equipment_type(ahu_id, df=frames[ahu_id], role_map=role_map)
        feeds = [v for v in children.get(ahu_id, []) if v in scope and v in frames]
        _add_heading(doc, f"{et} — {ahu_id}", 1)
        if feeds:
            _add_para(doc, "feeds (VAV children): " + ", ".join(feeds))
        _append_cards(ahu_id, et)
        feed_faults = [
            r
            for r in results
            if r.equipment_id in feeds
            and r.rule_id in FEED_RELATIONSHIP_RULE_IDS
            and str(r.status) in {"FAULT", "WARNING", "PASS", "SKIPPED_MISSING_ROLES"}
        ]
        if feed_faults or feeds:
            _add_heading(doc, f"Feed relationship faults (VAVs served by {ahu_id})", 2)
            _add_para(
                doc,
                "VAV leave vs parent AHU SAT and related topology-enabled checks.",
            )
            for vav_id in feeds:
                vav_et = resolve_equipment_type(vav_id, df=frames[vav_id], role_map=role_map)
                _add_heading(doc, f"fedBy {ahu_id} ← {vav_id}", 3)
                _append_cards(vav_id, vav_et, rule_ids=list(FEED_RELATIONSHIP_RULE_IDS))

    # --- One VAV / Zones chapter ---
    vav_ids = sorted(
        [
            eq
            for eq in scope
            if eq in frames
            and resolve_equipment_type(eq, df=frames[eq], role_map=role_map) == "VAV"
        ]
    )
    _add_heading(doc, "VAV / Zones — cohort analysis", 1)
    _add_para(
        doc,
        "Zone-level analytics and worst-performing VAV ranking (occupied comfort fails). "
        "Individual VAV fault cards stay in Plots; this chapter is the building-wide view.",
    )
    _add_para(doc, "Zone comfort ranking chart", bold=True)
    _add_para(doc, PLACE_RCX_PLOT_HERE.format(preset_id="zone_comfort_rank"))
    try:
        schedule = OccupancySchedule.from_dict(occupancy_schedule)
        rank = zone_comfort_fail_ranking(
            frames,
            role_map,
            schedule=schedule,
            comfort_low_f=zone_lo_f,
            comfort_high_f=zone_hi_f,
            equipment_types=("VAV",),
        )
        if vav_ids:
            rank = rank[rank["equipment_id"].isin(vav_ids)] if not rank.empty else rank
    except Exception:
        rank = pd.DataFrame()
    if rank is None or rank.empty:
        _add_para(doc, "[Placeholder] No zone comfort ranking rows for this package.")
    else:
        cols = [
            c
            for c in rank.columns
            if c
            in {
                "equipment_id",
                "pct_outside_comfort",
                "n_occupied",
                "n_outside",
                "outlier",
                "mean_zone_t",
            }
        ] or list(rank.columns)[:6]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = str(c)
        for _, row in rank.head(40).iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(cols):
                cells[i].text = str(row[c])
    # Compact VAV fault tally (not per-box cards)
    if vav_ids:
        _add_heading(doc, "VAV fault tally (run scope)", 2)
        rows = []
        for r in results:
            if r.equipment_id not in vav_ids:
                continue
            if r.rule_id in FEED_RELATIONSHIP_RULE_IDS:
                continue  # already under AHU
            if str(r.status) not in {"FAULT", "WARNING"}:
                continue
            rows.append((r.equipment_id, r.rule_id, r.status, getattr(r, "fault_hours", None)))
        if not rows:
            _add_para(doc, "No FAULT/WARNING VAV results in this run (excluding feed rules).")
        else:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for i, h in enumerate(["Equipment", "Rule", "Status", "Fault hours"]):
                table.rows[0].cells[i].text = h
            for eq_id, rid, st, fh in rows[:80]:
                cells = table.add_row().cells
                cells[0].text = str(eq_id)
                cells[1].text = str(rid)
                cells[2].text = str(st)
                cells[3].text = "—" if fh is None else f"{fh:.2f}"

    # --- Other plant / weather equipment in scope ---
    other_ids = sorted(
        [
            eq
            for eq in scope
            if eq in frames
            and resolve_equipment_type(eq, df=frames[eq], role_map=role_map)
            not in {"AHU", "RTU", "VAV"}
        ]
    )
    for eq_id in other_ids:
        et = resolve_equipment_type(eq_id, df=frames[eq_id], role_map=role_map)
        _add_heading(doc, f"{et} — {eq_id}", 1)
        _append_cards(eq_id, et)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_session_docx_pack(
    *,
    building_id: str,
    frames: dict[str, pd.DataFrame],
    role_map: dict,
    results: list[RuleResult],
    equipment_ids: list[str] | None = None,
    params: dict[str, Any] | None = None,
    weather: pd.DataFrame | None = None,
    vav_to_ahu: dict[str, str] | None = None,
    motor_weekly: pd.DataFrame | None = None,
    cool_bins: pd.DataFrame | None = None,
    rcx_coverage: pd.DataFrame | None = None,
    zone_lo_f: float = 70.0,
    zone_hi_f: float = 75.0,
    occupancy_schedule: dict | None = None,
    plant_pump_summaries: dict[str, pd.DataFrame] | None = None,
    fan_on_summaries: dict[str, pd.DataFrame] | None = None,
) -> bytes:
    """ZIP: fdd_by_system + analytics + rcx_catalog + data_model."""
    import zipfile

    from app.rcx_plots import rcx_preset_coverage

    scope_frames = {
        k: v
        for k, v in frames.items()
        if equipment_ids is None or k in set(equipment_ids)
    } or frames
    tree = build_data_model_tree(
        scope_frames,
        role_map,
        building_id=building_id or "",
        vav_to_ahu=vav_to_ahu,
    )
    cov = rcx_coverage
    if cov is None or (isinstance(cov, pd.DataFrame) and cov.empty):
        try:
            cov = rcx_preset_coverage(
                frames,
                role_map,
                weather=weather,
                comfort_low_f=zone_lo_f,
                comfort_high_f=zone_hi_f,
            )
        except Exception:
            cov = pd.DataFrame()

    fdd = build_fdd_by_system_docx(
        building_id=building_id,
        frames=frames,
        role_map=role_map,
        results=results,
        equipment_ids=equipment_ids,
        params=params,
        vav_to_ahu=vav_to_ahu,
        zone_lo_f=zone_lo_f,
        zone_hi_f=zone_hi_f,
        occupancy_schedule=occupancy_schedule,
    )
    analytics = build_analytics_docx(
        building_id=building_id,
        motor_weekly=motor_weekly,
        cool_bins=cool_bins,
        rcx_coverage=cov if isinstance(cov, pd.DataFrame) else None,
        tree=tree,
        plant_pump_summaries=plant_pump_summaries,
        fan_on_summaries=fan_on_summaries,
    )
    rcx = build_rcx_catalog_docx(
        building_id=building_id,
        frames=frames,
        role_map=role_map,
        weather=weather,
        results=results,
        params=params,
        rcx_coverage=cov if isinstance(cov, pd.DataFrame) else None,
        zone_lo_f=zone_lo_f,
        zone_hi_f=zone_hi_f,
        occupancy_schedule=occupancy_schedule,
        motor_weekly=motor_weekly,
        cool_bins=cool_bins,
    )
    data_model = build_building_data_model_docx(tree)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("fdd_by_system.docx", fdd)
        zf.writestr("analytics.docx", analytics)
        zf.writestr("rcx_catalog.docx", rcx)
        zf.writestr("data_model.docx", data_model)
    return buf.getvalue()
