#!/usr/bin/env python3
"""Controls-service checklist from vibe19 WattLab dumps (VAV / sensors / hunting).

Pure dump consumer — no EnergyPlus / utility scoring. Reads FDD + RCx comfort +
sensor health / fan-on vs fan-off stats and emits Markdown + JSON (+ optional DOCX).

Also flags unusually high / suspect false-positive FDD rows. Agents should iterate
vibe19 FDD tuning when positives look epidemic, then pass ``--fp-tuning-notes`` so
MD/DOCX reports record before/after tuning attempts.

Prefer the packaged CLI::

    wattlab controls-checklist --dump … --out-dir … --docx --fp-tuning-notes notes.md

Workspace script form (same module)::

    docker exec -i vibe20 python /data/tools/controls_service_checklist.py \\
      --dump /data/uploads/dump/wattlab_dump_BUILDING_100.zip \\
      --out-dir /data/reports/controls_checklist \\
      --docx
"""
from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path
from typing import Any

import pandas as pd

# ---------------------------------------------------------------------------
# Rule labels + soft ECM flags (edit here; no deep ECM writeup)
# ---------------------------------------------------------------------------

RULE_LABELS: dict[str, str] = {
    "VAV-1": "Zone outside comfort band (occupied)",
    "VAV-3": "Excessive reheat in warm weather",
    "VAV-4": "Damper stuck near full open",
    "VAV-5": "Airflow sensor bias (flow with damper closed)",
    "VAV-7": "Min / fixed high airflow",
    "VAV-REHEAT": "Reheat open with little DAT rise",
    "PID-HUNT-1": "Suspected AO hunting (damper/valve/fan)",
    "DMP-1": "OA damper leakage / stuck",
    "VLV-1": "Cooling valve leakage / stuck",
    "SV-FLATLINE": "Sensor flatline / stuck",
    "SV-SPIKE": "Sensor spike",
    "SV-STALE": "Sensor stale",
    "SV-RATE": "Implausible sensor rate of change",
    "SV-RANGE": "Sensor out of expected range",
}

ECM_SOFT_FLAGS: dict[str, str] = {
    "VAV-1": "Investigate zone comfort / box control / sensor",
    "VAV-3": "Check reheat valve + SAT / OA lockout",
    "VAV-4": "Investigate stuck damper / box airflow control",
    "VAV-5": "Airflow sensor bias / damper position feedback",
    "VAV-7": "Review min-flow SP / stuck damper at high flow",
    "VAV-REHEAT": "Check reheat valve stroke + DAT sensor",
    "PID-HUNT-1": "Tune PID / check AO wiring & feedback",
    "DMP-1": "OA damper actuator / seal / linkage",
    "VLV-1": "Cooling valve actuator / seat leakage",
    "SV-FLATLINE": "Replace or rescale sensor",
    "SV-SPIKE": "Check sensor wiring / noise / grounding",
    "SV-STALE": "Check BACnet poll / controller / point map",
    "SV-RATE": "Check sensor / scaling / bad spike filter",
    "SV-RANGE": "Verify sensor span / calibration",
}

VAV_MECH_RULES = frozenset({"VAV-3", "VAV-4", "VAV-5", "VAV-7", "VAV-REHEAT"})
SV_RULES = frozenset({"SV-FLATLINE", "SV-SPIKE", "SV-STALE", "SV-RATE", "SV-RANGE"})

# Fan-off smell: roles where a high reading with fan off is suspicious
FAN_OFF_PRESSURE_ROLES = frozenset(
    {
        "duct-static-pressure",
        "supply-fan-differential-pressure",
        "filter-differential-pressure",
        "building-static-pressure",
    }
)
FAN_OFF_STATIC_WC_THRESH = 1.0  # in. w.c. with fan off → flag

# Unusual / suspect high fault-positive heuristics (re-check in vibe19 FDD Plots)
HIGH_FAULT_PCT = 95.0
NEAR_WINDOW_FRAC = 0.85  # fault_hours / span_hours
EPIDEMIC_VAV_FRAC = 0.40  # same rule FAULT on ≥40% of VAVs

COLUMN_HELP = {
    "exec": (
        "**What these counts mean:** each row is a count of checklist hits from the vibe19 dump "
        "(not energy use). “Below in-band” uses the same occupied comfort band as vibe19 RCx Plots "
        "(default ~70–75°F): **in-band % = 100 − pct_outside_comfort**. "
        "`fault_hours` / `fault_pct` come from FDD (`fdd_summary`): hours (and % of *active/gated* "
        "samples) where the rule’s confirmed fault was true."
    ),
    "comfort": (
        "**Columns:** `in_band_%` = share of occupied samples inside the comfort band; "
        "`outside_%` = vibe19 `pct_outside_comfort`; `n_below` / `n_above` = too-cold / too-hot "
        "occupied samples; `mean_T`/`min`/`max` = zone-air-temp (°F). "
        "`dead_sensor?` is set when mean zone T is far below the band or the RCx `outlier` flag is set — "
        "often a failed or unscaled zone sensor, not a real comfort complaint."
    ),
    "vav_mech": (
        "**Columns:** OpenFDD VAV mechanical rules only. "
        "`VAV-4` = damper stuck near full open; `VAV-5` = airflow while damper nearly closed "
        "(flow sensor bias or bad damper feedback); `VAV-7` = min/fixed high airflow; "
        "`VAV-3` = reheat open in warm weather; `fault_pct` = % of *active* samples in fault "
        "(can be high even when hours are modest if the gate is narrow)."
    ),
    "sensors": (
        "**Fan-off anomalies:** compare `sensor_stats_fan_off` vs `fan_on` for duct/building static. "
        "A high static with the fan OFF (while ON looks normal) usually means a bad/stuck pressure "
        "sensor — classic Liberty AHU daytime-OK / night-wild pattern. "
        "**SV rows:** sensor-validation sweeps (flatline / spike / stale / rate / range) with hours "
        "and typical values (`p50`, `max`)."
    ),
    "other": (
        "**PID-HUNT-1** flags AO hunting on dampers/valves/fans when the rule FAULTs. "
        "**Other HVAC** is every remaining FDD FAULT (schedules, CHW, economizer, AHU, etc.). "
        "Near-100% `fault_pct` schedule/plant rules are often configuration or proof issues — "
        "open them in vibe19 FDD Plots before treating as field work."
    ),
    "ecm": (
        "**Soft flags only:** a one-line suggested investigation theme keyed by `rule_id`. "
        "Not a measure list, ROI, or calibrated ECM. Use to prioritize tech walk-downs."
    ),
    "unusual": (
        "**Suspect high / false-positive smells:** faults that fire for nearly the whole data "
        "window, hit ≥95% of active samples, appear as an epidemic across many VAVs, or conflict "
        "with a quick telemetry spot-check (e.g. VAV-5 while damper is not actually closed). "
        "Re-plot these in vibe19 (FDD Plots + RCx) before dispatching. When counts look "
        "excessively high, agents should iterate vibe19 FDD tuning (thresholds / gates / role map) "
        "and record attempts under **Agent FDD false-positive tuning** in this report."
    ),
    "gaps": (
        "Skipped VAV rules usually mean missing `role_map` points (damper, reheat-valve, flow, etc.). "
        "Fix mapping in vibe19 before trusting PASS/FAULT on those rules."
    ),
    "fp_tuning": (
        "**Agent duty:** if unusual/suspect rows or epidemic VAV rules dominate, do not publish "
        "the first dump as gospel. Iterate in vibe19 (retune FDD positives, fix role map, "
        "re-export dump), then regenerate this checklist and summarize what changed here."
    ),
}


def _span_hours(data_window: Any) -> float | None:
    if not isinstance(data_window, dict):
        return None
    sp = data_window.get("span_hours")
    try:
        return float(sp) if sp is not None else None
    except (TypeError, ValueError):
        return None


def telemetry_role_medians(zf: zipfile.ZipFile, equipment_id: str) -> dict[str, float]:
    """Median of common VAV/AHU roles from telemetry/<id>.csv if present."""
    path = f"telemetry/{equipment_id}.csv"
    try:
        df = pd.read_csv(zf.open(path))
    except KeyError:
        return {}
    want = (
        "damper",
        "reheat-valve",
        "zone-airflow",
        "zone-air-temp",
        "cooling-setpoint",
        "effective-setpoint",
        "duct-static-pressure",
        "fan-status",
        "discharge-air-temp",
    )
    out: dict[str, float] = {}
    cols = {c.lower(): c for c in df.columns}
    for role in want:
        col = cols.get(role) or (role if role in df.columns else None)
        if col is None:
            continue
        s = pd.to_numeric(df[col], errors="coerce").dropna()
        if len(s) >= 20:
            out[role] = round(float(s.median()), 3)
    return out


def analyze_unusual_faults(
    report: dict[str, Any],
    *,
    zf: zipfile.ZipFile | None,
    n_vav: int,
) -> dict[str, Any]:
    """Flag unusually high / suspect false-positive FDD rows; spot-check dump telemetry."""
    span = _span_hours(report["summary"].get("data_window"))
    faults = list(report["fdd"]["all_faults"])
    comfort_dead = {
        r["equipment_id"]
        for r in report["comfort"]["rows"]
        if r.get("flag_dead_sensor")
    }

    # Epidemic rules among VAV mechanical
    vav_rule_counts: dict[str, int] = {}
    for x in report["fdd"]["vav_mech"]:
        vav_rule_counts[x["rule_id"]] = vav_rule_counts.get(x["rule_id"], 0) + 1
    epidemic_rules = {
        rid
        for rid, n in vav_rule_counts.items()
        if n_vav > 0 and n / n_vav >= EPIDEMIC_VAV_FRAC
    }

    suspects: list[dict[str, Any]] = []
    for x in faults:
        reasons: list[str] = []
        hours = x["fault_hours"]
        pct = x["fault_pct"]
        rid = x["rule_id"]
        eid = x["equipment_id"]

        if pct >= HIGH_FAULT_PCT:
            reasons.append(f"fault_pct {pct:.1f}% ≥ {HIGH_FAULT_PCT}% of active samples")
        if span and hours >= NEAR_WINDOW_FRAC * span:
            reasons.append(
                f"fault_hours {hours:.0f} ≈ {100 * hours / span:.0f}% of data window "
                f"({span:.0f} h) — near-continuous"
            )
        if rid in epidemic_rules:
            reasons.append(
                f"{rid} FAULTs on {vav_rule_counts.get(rid, 0)}/{n_vav} VAVs "
                f"(≥{100 * EPIDEMIC_VAV_FRAC:.0f}% — possible mapping/proof epidemic)"
            )
        if eid in comfort_dead and rid.startswith("VAV"):
            reasons.append("same box flagged dead/outlier zone sensor in comfort ranking")
        if str(rid).upper().startswith("SCHED") and pct >= 90:
            reasons.append("schedule rule near-always faulting — check occupied/proof config in vibe19")

        tel: dict[str, float] = {}
        if reasons and zf is not None:
            tel = telemetry_role_medians(zf, eid)
            if rid == "VAV-5" and tel:
                d = tel.get("damper")
                flow = tel.get("zone-airflow")
                if d is not None and d > 25:
                    reasons.append(
                        f"telemetry median damper={d}% (not closed) — VAV-5 may be false positive / bad damper feedback"
                    )
                elif d is not None and d <= 10 and flow is not None and flow > 50:
                    reasons.append(
                        f"telemetry supports bias: damper≈{d}%, flow≈{flow} — likely real sensor issue"
                    )
            if rid == "VAV-4" and tel.get("damper") is not None:
                d = tel["damper"]
                if d >= 90:
                    reasons.append(f"telemetry median damper={d}% — consistent with stuck-open")
                elif d < 50:
                    reasons.append(
                        f"telemetry median damper={d}% (not full open) — re-check VAV-4 gate in vibe19"
                    )
            if rid == "VAV-1" and tel.get("zone-air-temp") is not None:
                zt = tel["zone-air-temp"]
                if zt < 40 or zt > 95:
                    reasons.append(f"telemetry median zone-air-temp={zt}°F — sensor/scale issue likely")

        if not reasons:
            continue
        severity = "high_suspect"
        if any("supports bias" in r or "consistent with stuck" in r for r in reasons):
            severity = "likely_real"
        elif any("false positive" in r or "epidemic" in r or "near-continuous" in r for r in reasons):
            severity = "high_suspect"
        suspects.append(
            {
                "equipment_id": eid,
                "rule_id": rid,
                "label": x.get("label") or rid,
                "fault_hours": hours,
                "fault_pct": pct,
                "severity": severity,
                "reasons": reasons,
                "telemetry_medians": tel,
                "vibe19_hint": (
                    f"In vibe19: FDD Plots → filter {rid} / {eid}; confirm swim-lane vs damper/flow/zone T"
                ),
            }
        )

    suspects.sort(
        key=lambda r: (
            0 if r["severity"] == "high_suspect" else 1,
            -r["fault_hours"],
        )
    )
    return {
        "n_suspects": len(suspects),
        "epidemic_vav_rules": sorted(epidemic_rules),
        "vav_rule_fault_counts": vav_rule_counts,
        "span_hours": span,
        "rows": suspects[:40],
        "agent_should_iterate_vibe19": bool(
            epidemic_rules
            or len([s for s in suspects if s["severity"] == "high_suspect"]) >= 3
            or (len(suspects) >= 8)
        ),
    }


def build_narratives(report: dict[str, Any]) -> dict[str, str]:
    """Paragraph-form read of the checklist for humans / DOCX."""
    s = report["summary"]
    bid = s["building_id"]
    span = _span_hours(s.get("data_window"))
    span_txt = f" about {span:.0f} hours of history" if span else " the available history"

    dead = [r for r in report["comfort"]["rows"] if r.get("flag_dead_sensor")]
    below = [r for r in report["comfort"]["rows"] if r.get("below_threshold")]
    top = s.get("top_faults_by_hours") or []
    top_txt = ", ".join(
        f"{x['equipment_id']} ({x['rule_id']}, {x['fault_hours']:.0f} h, {x['fault_pct']:.0f}% active)"
        for x in top[:5]
    ) or "none"

    fan = report.get("fan_off_anomalies") or []
    fan_txt = ""
    if fan:
        bits = [
            f"{x['equipment_id']} {x['role']} reads {x['fan_off_p50']} {x.get('units','')} with fan OFF"
            f" (fan ON ≈ {x.get('fan_on_p50')})"
            for x in fan[:3]
        ]
        fan_txt = (
            " Fan-off pressure smell test flags: "
            + "; ".join(bits)
            + ". That pattern (wild at night / OK by day) is typical of a bad static sensor, "
            "not a real duct pressure."
        )

    unusual = report.get("unusual_faults") or {}
    n_sus = unusual.get("n_suspects") or 0
    epi = unusual.get("epidemic_vav_rules") or []
    sus_rows = unusual.get("rows") or []
    high_sus = [r for r in sus_rows if r.get("severity") == "high_suspect"]

    exec_p = (
        f"{bid} has {s['n_vav']} VAV boxes in the dump across{span_txt}. "
        f"{s['n_comfort_below_threshold']} zones sit below the {s['in_band_min']}% in-band comfort "
        f"threshold (vibe19 RCx band), and {len(dead)} of those look like dead or wildly wrong "
        f"zone temperature sensors rather than ordinary comfort drift. "
        f"VAV mechanical FDD shows {s['n_vav_mech_faults']} fault rows "
        f"({s['n_damper_stuck_vav4']} damper-stuck VAV-4, {s['n_flow_bias_vav5']} flow-bias VAV-5). "
        f"Sensor validation adds {s['n_sv_fault_rows']} SV FDD rows and "
        f"{s['n_sensor_fault_rows']} sensor-fault summary rows; "
        f"PID hunting FAULTs: {s['n_pid_hunt_faults']}; other HVAC FAULTs: {s['n_other_hvac_faults']}. "
        f"Largest fault-hour hits: {top_txt}."
        f"{fan_txt}"
    )

    if below:
        names = ", ".join(
            f"{r['equipment_id']} ({r['in_band_pct']:.0f}% in-band, mean {r['mean_zone_t']}°F)"
            for r in below[:6]
        )
        comfort_p = (
            f"Worst comfort performers (lowest in-band %): {names}. "
            "Treat `dead_sensor?` / near-freezing or near-zero means as instrumentation first — "
            "tuning the box will not fix a failed zone sensor. "
            "Boxes with middling in-band % and lots of `n_above`/`n_below` are likelier true "
            "control or load problems; open them next to SAT and neighboring zones in vibe19 RCx."
        )
    else:
        comfort_p = (
            f"No VAVs fell below the {s['in_band_min']}% in-band threshold under the dump’s "
            "occupied comfort ranking."
        )

    mech = report["fdd"]["vav_mech"]
    if mech:
        by: dict[str, int] = {}
        for x in mech:
            by[x["rule_id"]] = by.get(x["rule_id"], 0) + 1
        by_txt = ", ".join(f"{k}×{v}" for k, v in sorted(by.items(), key=lambda kv: -kv[1]))
        worst = mech[0]
        mech_p = (
            f"VAV mechanical faults by rule: {by_txt}. "
            f"Highest hours: {worst['equipment_id']} / {worst['rule_id']} "
            f"({worst['fault_hours']:.0f} h, {worst['fault_pct']:.0f}% of active) — {worst['label']}. "
        )
        if epi:
            mech_p += (
                f"Epidemic rules (≥{100 * EPIDEMIC_VAV_FRAC:.0f}% of VAVs): {', '.join(epi)}. "
                "When many boxes fail the same rule, prefer a shared mapping, proof, or AHU "
                "issue over dozens of independent actuator failures."
            )
        else:
            mech_p += "No single VAV mechanical rule covers an epidemic share of boxes."
    else:
        mech_p = "No VAV-3/4/5/7/REHEAT FAULT rows in this dump."

    sens_p = (
        f"Sensor section lists {len(fan)} fan-off pressure anomal"
        f"{'y' if len(fan) == 1 else 'ies'} and the top SV/sensor-fault rows by hours. "
        "Flatlines on temps/pressures and stale BACnet points are usually wiring, scaling, or "
        "poll issues; confirm with vibe19 Sensor health before replacing hardware."
    )

    if n_sus:
        examples = "; ".join(
            f"{r['equipment_id']}/{r['rule_id']} ({r['severity']}: {r['reasons'][0]})"
            for r in high_sus[:4] or sus_rows[:4]
        )
        unusual_p = (
            f"Automated review flagged {n_sus} unusually high or suspect fault rows "
            f"({len(high_sus)} high-suspect). Examples: {examples}. "
            "For each, use the vibe19 hint (FDD Plots swim lane + related trends) to decide "
            "real failure vs false positive before writing work orders."
        )
        if unusual.get("agent_should_iterate_vibe19"):
            unusual_p += (
                " Fault positives look excessively high — agents should attempt vibe19 FDD "
                "retuning (thresholds/gates/role map), re-export the dump, and record the "
                "tuning attempt under Agent FDD false-positive tuning in this DOCX/MD report."
            )
    else:
        unusual_p = (
            "No faults met the automated near-window / epidemic / telemetry-conflict heuristics."
        )

    other_n = s["n_other_hvac_faults"]
    other_p = (
        f"There are {other_n} non-VAV-mechanical / non-SV FAULT rows (schedules, plants, AHUs, "
        "economizers, etc.). Prioritize those with both high hours and plausible field meaning; "
        "de-prioritize near-100% schedule or meteo-diff rules until proofs/mapping are checked."
    )

    return {
        "executive": exec_p,
        "comfort": comfort_p,
        "vav_mech": mech_p,
        "sensors": sens_p,
        "unusual": unusual_p,
        "other": other_p,
    }


def _slug(name: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")
    return s or "building"


def _read_csv(zf: zipfile.ZipFile, name: str) -> pd.DataFrame | None:
    try:
        with zf.open(name) as fh:
            return pd.read_csv(fh)
    except KeyError:
        return None


def _read_json(zf: zipfile.ZipFile, name: str) -> Any | None:
    try:
        with zf.open(name) as fh:
            return json.loads(fh.read().decode("utf-8", errors="replace"))
    except KeyError:
        return None


def _safe_float(v: Any, default: float = 0.0) -> float:
    try:
        if v is None or (isinstance(v, float) and pd.isna(v)):
            return default
        return float(v)
    except (TypeError, ValueError):
        return default


def _status_fault(s: Any) -> bool:
    return str(s).strip().upper() == "FAULT"


def load_dump(path: Path) -> dict[str, Any]:
    """Load checklist-relevant tables from a vibe19 dump zip."""
    zf = zipfile.ZipFile(path)
    seed = _read_json(zf, "model_seed.json") or {}
    building_id = (
        seed.get("display_name")
        or seed.get("project_id")
        or seed.get("building_id")
        or path.stem.replace("wattlab_dump_", "")
    )
    data_window = seed.get("data_window") or {}
    return {
        "path": str(path),
        "building_id": str(building_id),
        "slug": _slug(str(building_id)),
        "data_window": data_window,
        "seed": seed,
        "comfort": _read_csv(zf, "rcx_zone_comfort_ranking.csv"),
        "fdd": _read_csv(zf, "fdd_summary.csv"),
        "sensor_fault": _read_csv(zf, "sensor_fault_summary.csv"),
        "sensor_health": _read_csv(zf, "sensor_health_matrix.csv"),
        "fan_off": _read_csv(zf, "sensor_stats_fan_off.csv"),
        "fan_on": _read_csv(zf, "sensor_stats_fan_on.csv"),
        "gap": _read_csv(zf, "role_map_gap_report.csv"),
        "data_model": _read_csv(zf, "data_model.csv"),
        "tuning": _read_json(zf, "tuning_assistant_report.json"),
        "zf": zf,  # kept closed by caller via context — close after analyze
    }


def _vav_count(fdd: pd.DataFrame | None, comfort: pd.DataFrame | None) -> int:
    ids: set[str] = set()
    if comfort is not None and "equipment_id" in comfort.columns:
        ids.update(comfort["equipment_id"].astype(str))
    if fdd is not None and "equipment_type" in fdd.columns:
        m = fdd["equipment_type"].astype(str).str.upper() == "VAV"
        ids.update(fdd.loc[m, "equipment_id"].astype(str))
    return len(ids)


def analyze_comfort(comfort: pd.DataFrame | None, in_band_min: float) -> dict[str, Any]:
    if comfort is None or comfort.empty:
        return {"rows": [], "below_threshold": [], "n_vav": 0, "n_below": 0}
    df = comfort.copy()
    df["pct_outside_comfort"] = pd.to_numeric(df["pct_outside_comfort"], errors="coerce")
    df["in_band_pct"] = (100.0 - df["pct_outside_comfort"]).round(2)
    for c in ("n_occupied", "n_outside", "n_below", "n_above", "mean_zone_t", "min_zone_t", "max_zone_t"):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")
    rows = []
    for _, r in df.sort_values("in_band_pct", ascending=True).iterrows():
        mean_t = _safe_float(r.get("mean_zone_t"), float("nan"))
        low = _safe_float(r.get("comfort_low_f"), 70.0)
        outlier = bool(r.get("outlier")) if "outlier" in df.columns else False
        deadish = outlier or (pd.notna(mean_t) and mean_t < low - 20)
        rows.append(
            {
                "equipment_id": str(r["equipment_id"]),
                "in_band_pct": _safe_float(r["in_band_pct"]),
                "pct_outside_comfort": _safe_float(r["pct_outside_comfort"]),
                "n_occupied": int(_safe_float(r.get("n_occupied"))),
                "n_below": int(_safe_float(r.get("n_below"))),
                "n_above": int(_safe_float(r.get("n_above"))),
                "mean_zone_t": round(mean_t, 2) if pd.notna(mean_t) else None,
                "min_zone_t": round(_safe_float(r.get("min_zone_t")), 2),
                "max_zone_t": round(_safe_float(r.get("max_zone_t")), 2),
                "comfort_low_f": _safe_float(r.get("comfort_low_f"), 70.0),
                "comfort_high_f": _safe_float(r.get("comfort_high_f"), 75.0),
                "outlier": outlier,
                "flag_dead_sensor": deadish,
                "below_threshold": _safe_float(r["in_band_pct"]) < in_band_min,
            }
        )
    below = [x for x in rows if x["below_threshold"]]
    return {
        "rows": rows,
        "below_threshold": below,
        "n_vav": len(rows),
        "n_below": len(below),
        "in_band_min": in_band_min,
    }


def _clean_str(v: Any) -> str:
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return ""
    s = str(v).strip()
    return "" if s.lower() in {"nan", "none", "null"} else s


def analyze_fdd(fdd: pd.DataFrame | None) -> dict[str, Any]:
    empty = {
        "vav_mech": [],
        "vav_comfort_fdd": [],
        "sv_faults": [],
        "pid_hunt": [],
        "other_hvac": [],
        "all_faults": [],
        "by_rule_hours": {},
    }
    if fdd is None or fdd.empty:
        return empty
    df = fdd.copy()
    df["fault_hours"] = pd.to_numeric(df.get("fault_hours"), errors="coerce").fillna(0.0)
    df["fault_pct"] = pd.to_numeric(df.get("fault_pct"), errors="coerce").fillna(0.0)
    faults = df[df["status"].map(_status_fault)].copy()

    def _row(r: pd.Series) -> dict[str, Any]:
        rid = str(r["rule_id"])
        return {
            "equipment_id": str(r["equipment_id"]),
            "equipment_type": str(r.get("equipment_type", "")),
            "rule_id": rid,
            "label": RULE_LABELS.get(rid, rid),
            "fault_hours": round(_safe_float(r["fault_hours"]), 2),
            "fault_pct": round(_safe_float(r["fault_pct"]), 2),
            "ecm_flag": ECM_SOFT_FLAGS.get(rid),
            "missing_roles": _clean_str(r.get("missing_roles", "")),
            "notes": _clean_str(r.get("notes", "")),
        }

    all_faults = sorted((_row(r) for _, r in faults.iterrows()), key=lambda x: -x["fault_hours"])
    et = faults["equipment_type"].astype(str).str.upper() if "equipment_type" in faults.columns else pd.Series([""] * len(faults))
    is_vav = et == "VAV"
    rid = faults["rule_id"].astype(str)

    vav_mech = [_row(r) for _, r in faults[is_vav & rid.isin(VAV_MECH_RULES)].iterrows()]
    vav_mech.sort(key=lambda x: -x["fault_hours"])
    vav_comfort = [_row(r) for _, r in faults[is_vav & (rid == "VAV-1")].iterrows()]
    vav_comfort.sort(key=lambda x: -x["fault_hours"])
    sv = [_row(r) for _, r in faults[rid.isin(SV_RULES)].iterrows()]
    sv.sort(key=lambda x: -x["fault_hours"])
    pid = [_row(r) for _, r in faults[rid == "PID-HUNT-1"].iterrows()]
    pid.sort(key=lambda x: -x["fault_hours"])
    # Everything else: AHU dampers/valves, plant, etc. (not VAV-1/mech, not SV, not PID)
    other = [
        x
        for x in all_faults
        if x["rule_id"] not in SV_RULES
        and x["rule_id"] != "PID-HUNT-1"
        and not (
            x["equipment_type"].upper() == "VAV"
            and x["rule_id"] in (VAV_MECH_RULES | {"VAV-1"})
        )
    ]
    other.sort(key=lambda x: -x["fault_hours"])

    by_rule: dict[str, float] = {}
    for x in all_faults:
        by_rule[x["rule_id"]] = round(by_rule.get(x["rule_id"], 0.0) + x["fault_hours"], 2)

    return {
        "vav_mech": vav_mech,
        "vav_comfort_fdd": vav_comfort,
        "sv_faults": sv,
        "pid_hunt": pid,
        "other_hvac": other,
        "all_faults": all_faults,
        "by_rule_hours": dict(sorted(by_rule.items(), key=lambda kv: -kv[1])),
    }


def analyze_sensor_faults(sf: pd.DataFrame | None) -> list[dict[str, Any]]:
    if sf is None or sf.empty:
        return []
    df = sf.copy()
    df["fault_hours"] = pd.to_numeric(df.get("fault_hours"), errors="coerce").fillna(0.0)
    rows = []
    for _, r in df.sort_values("fault_hours", ascending=False).iterrows():
        hours = _safe_float(r["fault_hours"])
        if hours <= 0:
            continue
        rid = str(r.get("rule_id", ""))
        rows.append(
            {
                "equipment_id": str(r["equipment_id"]),
                "sensor": str(r.get("sensor", "")),
                "sensor_type": str(r.get("sensor_type", "")),
                "rule_id": rid,
                "label": RULE_LABELS.get(rid, rid),
                "fault_hours": round(hours, 2),
                "mean": round(_safe_float(r.get("mean")), 3) if pd.notna(r.get("mean")) else None,
                "p50": round(_safe_float(r.get("p50")), 3) if pd.notna(r.get("p50")) else None,
                "min": round(_safe_float(r.get("min")), 3) if pd.notna(r.get("min")) else None,
                "max": round(_safe_float(r.get("max")), 3) if pd.notna(r.get("max")) else None,
                "fault_max": round(_safe_float(r.get("fault_max")), 3)
                if pd.notna(r.get("fault_max"))
                else None,
                "ecm_flag": ECM_SOFT_FLAGS.get(rid),
            }
        )
    return rows


def analyze_fan_off_anomalies(
    fan_off: pd.DataFrame | None,
    fan_on: pd.DataFrame | None,
    *,
    static_thresh: float = FAN_OFF_STATIC_WC_THRESH,
) -> list[dict[str, Any]]:
    """Flag sensors that read absurdly high while the fan proof is off."""
    if fan_off is None or fan_off.empty or "role" not in fan_off.columns:
        return []
    off = fan_off.copy()
    on = fan_on.copy() if fan_on is not None and not fan_on.empty else None
    for c in ("mean", "p50", "median_fan_off", "median_fan_on"):
        if c in off.columns:
            off[c] = pd.to_numeric(off[c], errors="coerce")
        if on is not None and c in on.columns:
            on[c] = pd.to_numeric(on[c], errors="coerce")

    on_idx: dict[tuple[str, str], pd.Series] = {}
    if on is not None and "role" in on.columns:
        for _, r in on.iterrows():
            on_idx[(str(r["equipment_id"]), str(r["role"]))] = r

    anomalies = []
    for _, r in off.iterrows():
        role = str(r["role"])
        if role not in FAN_OFF_PRESSURE_ROLES and "static-pressure" not in role:
            continue
        if role.endswith("-sp") or role.endswith("setpoint"):
            continue
        off_med = r.get("p50")
        if pd.isna(off_med):
            off_med = r.get("mean")
        off_med = _safe_float(off_med, float("nan"))
        if pd.isna(off_med) or off_med < static_thresh:
            continue
        key = (str(r["equipment_id"]), role)
        on_row = on_idx.get(key)
        on_med = None
        if on_row is not None:
            on_med = on_row.get("p50")
            if pd.isna(on_med):
                on_med = on_row.get("mean")
            on_med = round(_safe_float(on_med), 3) if pd.notna(on_med) else None
        units = str(r.get("units", "") or "")
        anomalies.append(
            {
                "equipment_id": str(r["equipment_id"]),
                "role": role,
                "units": units,
                "fan_off_p50": round(off_med, 3),
                "fan_on_p50": on_med,
                "threshold": static_thresh,
                "note": (
                    f"Reads {off_med:.2f} {units} with fan OFF "
                    f"(threshold {static_thresh}); likely bad sensor / stuck high"
                ),
                "ecm_flag": "Replace or rescale duct static / pressure sensor",
            }
        )
    anomalies.sort(key=lambda x: -x["fan_off_p50"])
    return anomalies


def analyze_gaps(gap: pd.DataFrame | None, fdd: pd.DataFrame | None) -> dict[str, Any]:
    rows: list[dict[str, Any]] = []
    if gap is not None and not gap.empty:
        cols = [c for c in gap.columns]
        for _, r in gap.head(80).iterrows():
            rows.append({c: (None if pd.isna(r[c]) else str(r[c])) for c in cols})
    skipped = []
    if fdd is not None and not fdd.empty and "status" in fdd.columns:
        m = fdd["status"].astype(str).str.upper().str.startswith("SKIPPED")
        sub = fdd[m]
        if "equipment_type" in sub.columns:
            sub = sub[sub["equipment_type"].astype(str).str.upper() == "VAV"]
        for _, r in sub.head(60).iterrows():
            skipped.append(
                {
                    "equipment_id": str(r["equipment_id"]),
                    "rule_id": str(r["rule_id"]),
                    "status": str(r["status"]),
                    "missing_roles": _clean_str(r.get("missing_roles", "")),
                }
            )
    return {"gap_report_sample": rows[:40], "vav_skipped_rules": skipped}


def ecm_soft_summary(report: dict[str, Any]) -> list[dict[str, Any]]:
    """Unique soft ECM flags with supporting fault hours."""
    bag: dict[str, dict[str, Any]] = {}

    def _add(rule_id: str, equipment_id: str, hours: float, flag: str | None) -> None:
        if not flag:
            return
        key = f"{rule_id}|{flag}"
        if key not in bag:
            bag[key] = {
                "rule_id": rule_id,
                "ecm_flag": flag,
                "fault_hours_total": 0.0,
                "equipment_ids": [],
                "n_equipment": 0,
            }
        bag[key]["fault_hours_total"] = round(bag[key]["fault_hours_total"] + hours, 2)
        if equipment_id not in bag[key]["equipment_ids"]:
            bag[key]["equipment_ids"].append(equipment_id)
            bag[key]["n_equipment"] = len(bag[key]["equipment_ids"])

    for x in report["fdd"]["all_faults"]:
        _add(x["rule_id"], x["equipment_id"], x["fault_hours"], x.get("ecm_flag"))
    for x in report["sensor_faults"]:
        _add(x["rule_id"], x["equipment_id"], x["fault_hours"], x.get("ecm_flag"))
    for x in report["fan_off_anomalies"]:
        _add("FAN-OFF-STATIC", x["equipment_id"], 0.0, x.get("ecm_flag"))

    out = list(bag.values())
    out.sort(key=lambda x: (-x["fault_hours_total"], -x["n_equipment"]))
    for o in out:
        o["equipment_ids"] = o["equipment_ids"][:25]  # trim
    return out


def build_report(
    dump: dict[str, Any],
    *,
    in_band_min: float,
    static_thresh: float = FAN_OFF_STATIC_WC_THRESH,
    fp_tuning_notes: str | None = None,
) -> dict[str, Any]:
    comfort = analyze_comfort(dump["comfort"], in_band_min)
    fdd = analyze_fdd(dump["fdd"])
    sensor_faults = analyze_sensor_faults(dump["sensor_fault"])
    fan_off = analyze_fan_off_anomalies(
        dump["fan_off"], dump["fan_on"], static_thresh=static_thresh
    )
    gaps = analyze_gaps(dump["gap"], dump["fdd"])
    n_vav = comfort["n_vav"] or _vav_count(dump["fdd"], dump["comfort"])

    top5 = fdd["all_faults"][:5]
    summary = {
        "building_id": dump["building_id"],
        "dump_path": dump["path"],
        "data_window": dump["data_window"],
        "n_vav": n_vav,
        "n_comfort_below_threshold": comfort["n_below"],
        "in_band_min": in_band_min,
        "n_vav_mech_faults": len(fdd["vav_mech"]),
        "n_damper_stuck_vav4": sum(1 for x in fdd["vav_mech"] if x["rule_id"] == "VAV-4"),
        "n_flow_bias_vav5": sum(1 for x in fdd["vav_mech"] if x["rule_id"] == "VAV-5"),
        "n_sv_fault_rows": len(fdd["sv_faults"]),
        "n_sensor_fault_rows": len(sensor_faults),
        "n_fan_off_anomalies": len(fan_off),
        "n_pid_hunt_faults": len(fdd["pid_hunt"]),
        "n_other_hvac_faults": len(fdd["other_hvac"]),
        "top_faults_by_hours": top5,
        "fault_hours_by_rule": fdd["by_rule_hours"],
    }
    report: dict[str, Any] = {
        "summary": summary,
        "comfort": comfort,
        "fdd": fdd,
        "sensor_faults": sensor_faults[:80],
        "fan_off_anomalies": fan_off,
        "gaps": gaps,
        "tuning_notes": None,
        "fp_tuning_notes": (fp_tuning_notes or "").strip() or None,
    }
    if dump.get("tuning"):
        t = dump["tuning"]
        if isinstance(t, dict):
            report["tuning_notes"] = {k: t[k] for k in list(t)[:8]}
        else:
            report["tuning_notes"] = str(t)[:500]
    report["ecm_soft_flags"] = ecm_soft_summary(report)
    report["unusual_faults"] = analyze_unusual_faults(
        report, zf=dump.get("zf"), n_vav=n_vav
    )
    report["narratives"] = build_narratives(report)
    report["column_help"] = COLUMN_HELP
    return report


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------


def _md_table(headers: list[str], rows: list[list[Any]]) -> str:
    if not rows:
        return "_(none)_\n"
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join("" if v is None else str(v) for v in row) + " |")
    return "\n".join(lines) + "\n"


def _render_fp_tuning_section_md(report: dict[str, Any]) -> list[str]:
    help_ = report.get("column_help") or COLUMN_HELP
    unusual = report.get("unusual_faults") or {}
    parts = ["\n## Agent FDD false-positive tuning\n"]
    parts.append(help_.get("fp_tuning", COLUMN_HELP["fp_tuning"]) + "\n\n")
    if unusual.get("agent_should_iterate_vibe19"):
        parts.append(
            "**Automated flag:** fault positives look excessively high "
            f"({unusual.get('n_suspects', 0)} unusual rows"
            + (
                f"; epidemic VAV rules: {', '.join(unusual.get('epidemic_vav_rules') or [])}"
                if unusual.get("epidemic_vav_rules")
                else ""
            )
            + "). Iterate vibe19 FDD before treating as a field punch list.\n\n"
        )
    notes = report.get("fp_tuning_notes")
    if notes:
        parts.append("### Agent tuning log\n\n")
        parts.append(notes.strip() + "\n")
    elif unusual.get("agent_should_iterate_vibe19"):
        parts.append(
            "_No `--fp-tuning-notes` supplied yet. After retuning vibe19 and re-exporting, "
            "re-run with `--fp-tuning-notes` summarizing thresholds/gates/role-map changes "
            "and before/after fault counts._\n"
        )
    else:
        parts.append(
            "_No agent FP-tuning log required for this dump (heuristics did not flag an epidemic). "
            "Still re-check individual high_suspect rows in vibe19 FDD Plots._\n"
        )
    dump_tuning = report.get("tuning_notes")
    if dump_tuning:
        parts.append("\n### Dump `tuning_assistant_report.json` (excerpt)\n\n")
        parts.append(f"```json\n{json.dumps(dump_tuning, indent=2, default=str)[:2000]}\n```\n")
    return parts


def render_markdown(report: dict[str, Any]) -> str:
    s = report["summary"]
    bid = s["building_id"]
    narr = report.get("narratives") or {}
    help_ = report.get("column_help") or COLUMN_HELP
    parts: list[str] = []
    parts.append(f"# Controls service checklist — {bid}\n")
    parts.append(
        f"- Dump: `{s['dump_path']}`\n"
        f"- Data window: `{json.dumps(s.get('data_window') or {}, separators=(',', ':'))}`\n"
        f"- VAV count: **{s['n_vav']}**\n"
        f"- Comfort in-band threshold: **{s['in_band_min']}%** "
        f"(same band as vibe19 RCx; in-band % = 100 − pct_outside_comfort)\n"
    )

    parts.append("\n## How to read this report\n")
    parts.append(
        "Tables are vibe19 dump aggregates (FDD + RCx + sensor stats). "
        "`fault_hours` = hours the rule’s confirmed fault was true; "
        "`fault_pct` = that time as a percent of *active/gated* samples (not always wall-clock). "
        "Comfort **in-band %** matches the vibe19 RCx donut “In band” slice. "
        "Use the **Unusual / suspect high faults** section before treating near-100% hits as field failures — "
        "re-open those in vibe19 FDD Plots. When positives look epidemic, iterate vibe19 FDD "
        "tuning and record attempts under **Agent FDD false-positive tuning**.\n"
    )

    parts.append("\n## Executive summary\n")
    if narr.get("executive"):
        parts.append(narr["executive"] + "\n")
    parts.append("\n" + help_["exec"] + "\n\n")
    parts.append(
        f"| Metric | Count |\n| --- | --- |\n"
        f"| VAVs below in-band {s['in_band_min']}% | {s['n_comfort_below_threshold']} |\n"
        f"| VAV mechanical FDD faults (VAV-3/4/5/7/REHEAT) | {s['n_vav_mech_faults']} |\n"
        f"| Damper stuck (VAV-4) | {s['n_damper_stuck_vav4']} |\n"
        f"| Flow bias (VAV-5) | {s['n_flow_bias_vav5']} |\n"
        f"| Sensor SV FDD fault rows | {s['n_sv_fault_rows']} |\n"
        f"| Sensor fault summary rows | {s['n_sensor_fault_rows']} |\n"
        f"| Fan-off pressure anomalies | {s['n_fan_off_anomalies']} |\n"
        f"| PID hunting FAULTs | {s['n_pid_hunt_faults']} |\n"
        f"| Other HVAC FAULTs | {s['n_other_hvac_faults']} |\n"
        f"| Unusual / suspect fault rows | {(report.get('unusual_faults') or {}).get('n_suspects', 0)} |\n"
    )
    parts.append("\n### Top faults by hours\n")
    parts.append(
        _md_table(
            ["equipment_id", "rule_id", "label", "fault_hours", "fault_pct"],
            [
                [x["equipment_id"], x["rule_id"], x["label"], x["fault_hours"], x["fault_pct"]]
                for x in s["top_faults_by_hours"]
            ],
        )
    )

    parts.extend(_render_fp_tuning_section_md(report))

    parts.append("\n## Unusual / suspect high faults (re-check in vibe19)\n")
    if narr.get("unusual"):
        parts.append(narr["unusual"] + "\n")
    parts.append("\n" + help_["unusual"] + "\n\n")
    unusual = report.get("unusual_faults") or {}
    if unusual.get("epidemic_vav_rules"):
        parts.append(
            f"Epidemic VAV rules: **{', '.join(unusual['epidemic_vav_rules'])}** "
            f"(counts: `{json.dumps(unusual.get('vav_rule_fault_counts') or {})}`)\n\n"
        )
    parts.append(
        _md_table(
            ["severity", "equipment_id", "rule_id", "hours", "pct", "reasons", "telemetry", "vibe19"],
            [
                [
                    x["severity"],
                    x["equipment_id"],
                    x["rule_id"],
                    x["fault_hours"],
                    x["fault_pct"],
                    "; ".join(x["reasons"][:3]),
                    ", ".join(f"{k}={v}" for k, v in (x.get("telemetry_medians") or {}).items())[:80],
                    x.get("vibe19_hint", "")[:90],
                ]
                for x in (unusual.get("rows") or [])[:25]
            ],
        )
    )

    parts.append("\n## Zone comfort (VAV) — worst in-band first\n")
    if narr.get("comfort"):
        parts.append(narr["comfort"] + "\n")
    parts.append("\n" + help_["comfort"] + "\n\n")
    comfort_rows = report["comfort"]["rows"]
    show = [r for r in comfort_rows if r["below_threshold"] or r["flag_dead_sensor"]] or comfort_rows[:15]
    parts.append(
        _md_table(
            [
                "equipment_id",
                "in_band_%",
                "outside_%",
                "n_below",
                "n_above",
                "mean_T",
                "min",
                "max",
                "flags",
            ],
            [
                [
                    r["equipment_id"],
                    r["in_band_pct"],
                    r["pct_outside_comfort"],
                    r["n_below"],
                    r["n_above"],
                    r["mean_zone_t"],
                    r["min_zone_t"],
                    r["max_zone_t"],
                    ",".join(
                        [
                            x
                            for x, ok in (
                                ("below_threshold", r["below_threshold"]),
                                ("dead_sensor?", r["flag_dead_sensor"]),
                                ("outlier", r["outlier"]),
                            )
                            if ok
                        ]
                    ),
                ]
                for r in show
            ],
        )
    )

    parts.append("\n## VAV box mechanical\n")
    if narr.get("vav_mech"):
        parts.append(narr["vav_mech"] + "\n")
    parts.append("\n" + help_["vav_mech"] + "\n\n")
    parts.append(
        _md_table(
            ["equipment_id", "rule_id", "label", "fault_hours", "fault_pct", "ecm_flag"],
            [
                [
                    x["equipment_id"],
                    x["rule_id"],
                    x["label"],
                    x["fault_hours"],
                    x["fault_pct"],
                    x.get("ecm_flag") or "",
                ]
                for x in report["fdd"]["vav_mech"]
            ],
        )
    )

    parts.append("\n## Sensors not reading right\n")
    if narr.get("sensors"):
        parts.append(narr["sensors"] + "\n")
    parts.append("\n" + help_["sensors"] + "\n")
    parts.append("\n### Fan-off pressure anomalies\n")
    parts.append(
        _md_table(
            ["equipment_id", "role", "fan_off_p50", "fan_on_p50", "units", "note"],
            [
                [
                    x["equipment_id"],
                    x["role"],
                    x["fan_off_p50"],
                    x["fan_on_p50"],
                    x["units"],
                    x["note"],
                ]
                for x in report["fan_off_anomalies"]
            ],
        )
    )
    parts.append("\n### SV / sensor fault summary (top by hours)\n")
    parts.append(
        _md_table(
            ["equipment_id", "sensor", "rule_id", "fault_hours", "p50", "max", "ecm_flag"],
            [
                [
                    x["equipment_id"],
                    x["sensor"],
                    x["rule_id"],
                    x["fault_hours"],
                    x["p50"],
                    x["max"],
                    x.get("ecm_flag") or "",
                ]
                for x in report["sensor_faults"][:40]
            ],
        )
    )

    parts.append("\n## Hunting / other HVAC\n")
    if narr.get("other"):
        parts.append(narr["other"] + "\n")
    parts.append("\n" + help_["other"] + "\n")
    parts.append("\n### PID hunting (PID-HUNT-1 FAULT)\n")
    parts.append(
        _md_table(
            ["equipment_id", "fault_hours", "fault_pct", "ecm_flag"],
            [
                [x["equipment_id"], x["fault_hours"], x["fault_pct"], x.get("ecm_flag") or ""]
                for x in report["fdd"]["pid_hunt"]
            ],
        )
    )
    parts.append("\n### Other HVAC FAULTs (non VAV-mech / non SV)\n")
    parts.append(
        _md_table(
            ["equipment_id", "type", "rule_id", "label", "fault_hours", "fault_pct"],
            [
                [
                    x["equipment_id"],
                    x["equipment_type"],
                    x["rule_id"],
                    x["label"],
                    x["fault_hours"],
                    x["fault_pct"],
                ]
                for x in report["fdd"]["other_hvac"][:40]
            ],
        )
    )

    parts.append("\n## ECM soft flags (no detail)\n")
    parts.append(help_["ecm"] + "\n\n")
    parts.append(
        _md_table(
            ["rule_id", "ecm_flag", "n_equipment", "fault_hours_total", "sample_equipment"],
            [
                [
                    x["rule_id"],
                    x["ecm_flag"],
                    x["n_equipment"],
                    x["fault_hours_total"],
                    ", ".join(x["equipment_ids"][:5]),
                ]
                for x in report["ecm_soft_flags"]
            ],
        )
    )

    parts.append("\n## Gaps / skipped VAV rules\n")
    parts.append(help_["gaps"] + "\n\n")
    skipped = report["gaps"].get("vav_skipped_rules") or []
    miss: dict[str, int] = {}
    for row in skipped:
        mr = (row.get("missing_roles") or "").strip()
        if not mr:
            continue
        for piece in re.split(r"[,;|]", mr):
            piece = piece.strip()
            if piece:
                miss[piece] = miss.get(piece, 0) + 1
    if miss:
        parts.append("Missing roles on skipped VAV rules (counts):\n")
        parts.append(
            _md_table(
                ["missing_role", "n_skipped_rows"],
                [[k, v] for k, v in sorted(miss.items(), key=lambda kv: -kv[1])[:30]],
            )
        )
    else:
        parts.append("_(no missing_roles listed on skipped VAV FDD rows)_\n")
    parts.append("\nSample skipped VAV rules:\n")
    parts.append(
        _md_table(
            ["equipment_id", "rule_id", "status", "missing_roles"],
            [
                [x["equipment_id"], x["rule_id"], x["status"], x["missing_roles"][:80]]
                for x in skipped[:25]
            ],
        )
    )

    parts.append(
        "\n---\n_Generated by `wattlab controls-checklist` "
        "(`wattlab.existing_building.controls_checklist`) — controls checklist only; "
        "not an energy calibration report._\n"
    )
    return "\n".join(parts)


def _add_fp_tuning_docx(doc: Any, report: dict[str, Any]) -> None:
    help_ = report.get("column_help") or COLUMN_HELP
    unusual = report.get("unusual_faults") or {}
    doc.add_heading("Agent FDD false-positive tuning", level=1)
    doc.add_paragraph(help_.get("fp_tuning", COLUMN_HELP["fp_tuning"]).replace("**", ""))
    if unusual.get("agent_should_iterate_vibe19"):
        epi = unusual.get("epidemic_vav_rules") or []
        doc.add_paragraph(
            "Automated flag: fault positives look excessively high "
            f"({unusual.get('n_suspects', 0)} unusual rows"
            + (f"; epidemic VAV rules: {', '.join(epi)}" if epi else "")
            + "). Iterate vibe19 FDD before treating as a field punch list."
        )
    notes = report.get("fp_tuning_notes")
    if notes:
        doc.add_heading("Agent tuning log", level=2)
        doc.add_paragraph(notes.strip())
    elif unusual.get("agent_should_iterate_vibe19"):
        doc.add_paragraph(
            "No --fp-tuning-notes supplied yet. After retuning vibe19 and re-exporting, "
            "re-run with --fp-tuning-notes summarizing thresholds/gates/role-map changes "
            "and before/after fault counts."
        )
    dump_tuning = report.get("tuning_notes")
    if dump_tuning:
        doc.add_heading("Dump tuning_assistant_report.json (excerpt)", level=2)
        doc.add_paragraph(json.dumps(dump_tuning, indent=2, default=str)[:1500])


def render_docx(report: dict[str, Any], path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    s = report["summary"]
    narr = report.get("narratives") or {}
    help_ = report.get("column_help") or COLUMN_HELP
    doc = Document()
    doc.add_heading(f"Controls service checklist — {s['building_id']}", level=0)
    p = doc.add_paragraph()
    p.add_run(
        f"Dump: {s['dump_path']}\n"
        f"VAVs: {s['n_vav']} | Comfort threshold: {s['in_band_min']}% in-band\n"
    ).font.size = Pt(10)

    doc.add_heading("How to read this report", level=1)
    doc.add_paragraph(
        "Tables are vibe19 dump aggregates. fault_hours = confirmed-fault hours; "
        "fault_pct = percent of active/gated samples. Comfort in-band % matches the "
        "vibe19 RCx donut. Re-check unusual/suspect highs in vibe19 FDD Plots before dispatch. "
        "When positives look epidemic, iterate vibe19 FDD tuning and record attempts in "
        "Agent FDD false-positive tuning."
    )

    doc.add_heading("Executive summary", level=1)
    if narr.get("executive"):
        doc.add_paragraph(narr["executive"])
    doc.add_paragraph(help_["exec"].replace("**", ""))
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Count"
    for label, val in [
        (f"VAVs below in-band {s['in_band_min']}%", s["n_comfort_below_threshold"]),
        ("VAV mechanical FDD faults", s["n_vav_mech_faults"]),
        ("Damper stuck (VAV-4)", s["n_damper_stuck_vav4"]),
        ("Flow bias (VAV-5)", s["n_flow_bias_vav5"]),
        ("Fan-off pressure anomalies", s["n_fan_off_anomalies"]),
        ("PID hunting FAULTs", s["n_pid_hunt_faults"]),
        ("Other HVAC FAULTs", s["n_other_hvac_faults"]),
        (
            "Unusual / suspect fault rows",
            (report.get("unusual_faults") or {}).get("n_suspects", 0),
        ),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val)

    doc.add_heading("Top faults by hours", level=2)
    t2 = doc.add_table(rows=1, cols=5)
    t2.style = "Table Grid"
    for i, h in enumerate(["equipment_id", "rule_id", "label", "hours", "pct"]):
        t2.rows[0].cells[i].text = h
    for x in s["top_faults_by_hours"]:
        cells = t2.add_row().cells
        cells[0].text = x["equipment_id"]
        cells[1].text = x["rule_id"]
        cells[2].text = x["label"]
        cells[3].text = str(x["fault_hours"])
        cells[4].text = str(x["fault_pct"])

    _add_fp_tuning_docx(doc, report)

    doc.add_heading("Unusual / suspect high faults", level=1)
    if narr.get("unusual"):
        doc.add_paragraph(narr["unusual"])
    doc.add_paragraph(help_["unusual"].replace("**", ""))
    unusual_rows = (report.get("unusual_faults") or {}).get("rows") or []
    if unusual_rows:
        tu = doc.add_table(rows=1, cols=5)
        tu.style = "Table Grid"
        for i, h in enumerate(["severity", "equipment_id", "rule_id", "hours", "reasons"]):
            tu.rows[0].cells[i].text = h
        for x in unusual_rows[:20]:
            cells = tu.add_row().cells
            cells[0].text = x["severity"]
            cells[1].text = x["equipment_id"]
            cells[2].text = x["rule_id"]
            cells[3].text = str(x["fault_hours"])
            cells[4].text = "; ".join(x["reasons"][:2])[:200]

    doc.add_heading("Zone comfort (below threshold / dead sensor)", level=1)
    if narr.get("comfort"):
        doc.add_paragraph(narr["comfort"])
    doc.add_paragraph(help_["comfort"].replace("**", ""))
    show = [
        r
        for r in report["comfort"]["rows"]
        if r["below_threshold"] or r["flag_dead_sensor"]
    ] or report["comfort"]["rows"][:15]
    t3 = doc.add_table(rows=1, cols=6)
    t3.style = "Table Grid"
    for i, h in enumerate(["equipment_id", "in_band_%", "outside_%", "mean_T", "n_below", "n_above"]):
        t3.rows[0].cells[i].text = h
    for r in show:
        cells = t3.add_row().cells
        cells[0].text = r["equipment_id"]
        cells[1].text = str(r["in_band_pct"])
        cells[2].text = str(r["pct_outside_comfort"])
        cells[3].text = str(r["mean_zone_t"])
        cells[4].text = str(r["n_below"])
        cells[5].text = str(r["n_above"])

    doc.add_heading("VAV box mechanical", level=1)
    if narr.get("vav_mech"):
        doc.add_paragraph(narr["vav_mech"])
    doc.add_paragraph(help_["vav_mech"].replace("**", ""))
    t4 = doc.add_table(rows=1, cols=5)
    t4.style = "Table Grid"
    for i, h in enumerate(["equipment_id", "rule_id", "hours", "pct", "ecm_flag"]):
        t4.rows[0].cells[i].text = h
    for x in report["fdd"]["vav_mech"][:40]:
        cells = t4.add_row().cells
        cells[0].text = x["equipment_id"]
        cells[1].text = x["rule_id"]
        cells[2].text = str(x["fault_hours"])
        cells[3].text = str(x["fault_pct"])
        cells[4].text = x.get("ecm_flag") or ""

    doc.add_heading("Fan-off pressure anomalies", level=1)
    if narr.get("sensors"):
        doc.add_paragraph(narr["sensors"])
    doc.add_paragraph(help_["sensors"].replace("**", ""))
    if report["fan_off_anomalies"]:
        t5 = doc.add_table(rows=1, cols=5)
        t5.style = "Table Grid"
        for i, h in enumerate(["equipment_id", "role", "fan_off_p50", "fan_on_p50", "note"]):
            t5.rows[0].cells[i].text = h
        for x in report["fan_off_anomalies"]:
            cells = t5.add_row().cells
            cells[0].text = x["equipment_id"]
            cells[1].text = x["role"]
            cells[2].text = str(x["fan_off_p50"])
            cells[3].text = str(x["fan_on_p50"])
            cells[4].text = x["note"][:120]
    else:
        doc.add_paragraph("(none)")

    doc.add_heading("Other HVAC", level=1)
    if narr.get("other"):
        doc.add_paragraph(narr["other"])

    doc.add_heading("ECM soft flags", level=1)
    doc.add_paragraph(help_["ecm"].replace("**", ""))
    t6 = doc.add_table(rows=1, cols=4)
    t6.style = "Table Grid"
    for i, h in enumerate(["rule_id", "ecm_flag", "n_equip", "hours"]):
        t6.rows[0].cells[i].text = h
    for x in report["ecm_soft_flags"]:
        cells = t6.add_row().cells
        cells[0].text = x["rule_id"]
        cells[1].text = x["ecm_flag"]
        cells[2].text = str(x["n_equipment"])
        cells[3].text = str(x["fault_hours_total"])

    doc.add_paragraph(
        "Generated by wattlab controls-checklist — controls checklist only."
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def render_campus_summary(reports: list[dict[str, Any]]) -> str:
    lines = ["# Campus controls checklist summary\n"]
    lines.append(
        _md_table(
            [
                "building",
                "n_vav",
                "comfort_below",
                "vav_mech",
                "vav4",
                "vav5",
                "fan_off_anom",
                "pid_hunt",
                "unusual",
                "fp_iterate?",
            ],
            [
                [
                    r["summary"]["building_id"],
                    r["summary"]["n_vav"],
                    r["summary"]["n_comfort_below_threshold"],
                    r["summary"]["n_vav_mech_faults"],
                    r["summary"]["n_damper_stuck_vav4"],
                    r["summary"]["n_flow_bias_vav5"],
                    r["summary"]["n_fan_off_anomalies"],
                    r["summary"]["n_pid_hunt_faults"],
                    (r.get("unusual_faults") or {}).get("n_suspects", 0),
                    "yes" if (r.get("unusual_faults") or {}).get("agent_should_iterate_vibe19") else "no",
                ]
                for r in reports
            ],
        )
    )
    lines.append("\n## Per-building top fault\n")
    for r in reports:
        top = (r["summary"].get("top_faults_by_hours") or [None])[0]
        if top:
            lines.append(
                f"- **{r['summary']['building_id']}**: "
                f"{top['equipment_id']} / {top['rule_id']} "
                f"({top['fault_hours']} h) — {top['label']}\n"
            )
    return "\n".join(lines)


def _load_fp_tuning_notes(path: str | None, inline: str | None) -> str | None:
    chunks: list[str] = []
    if inline and inline.strip():
        chunks.append(inline.strip())
    if path:
        p = Path(path)
        if not p.is_file():
            raise SystemExit(f"fp-tuning-notes file not found: {p}")
        chunks.append(p.read_text(encoding="utf-8").strip())
    if not chunks:
        return None
    return "\n\n".join(chunks)


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument(
        "--dump",
        action="append",
        required=True,
        help="Path to vibe19 wattlab_dump_*.zip (repeatable)",
    )
    p.add_argument(
        "--out-dir",
        default="reports/controls_checklist",
        help="Output directory for md/json/docx",
    )
    p.add_argument(
        "--in-band-min",
        type=float,
        default=80.0,
        help="Flag VAVs with in-band %% below this (default 80)",
    )
    p.add_argument(
        "--docx",
        action="store_true",
        help="Also write .docx (requires python-docx)",
    )
    p.add_argument(
        "--static-fan-off-thresh",
        type=float,
        default=FAN_OFF_STATIC_WC_THRESH,
        help="Flag duct static (in. w.c.) at/above this with fan OFF",
    )
    p.add_argument(
        "--fp-tuning-notes",
        default=None,
        help="Path to markdown/text log of vibe19 FDD FP-tuning attempts (embedded in MD/DOCX)",
    )
    p.add_argument(
        "--fp-tuning-note",
        default=None,
        help="Inline FP-tuning note (combined with --fp-tuning-notes if both set)",
    )
    args = p.parse_args(argv)

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    static_thresh = float(args.static_fan_off_thresh)
    fp_notes = _load_fp_tuning_notes(args.fp_tuning_notes, args.fp_tuning_note)

    reports: list[dict[str, Any]] = []
    written: list[dict[str, str]] = []
    docx_ok = True
    docx_err: str | None = None

    for dump_s in args.dump:
        dump_path = Path(dump_s)
        if not dump_path.is_file():
            raise SystemExit(f"Dump not found: {dump_path}")
        dump = load_dump(dump_path)
        try:
            report = build_report(
                dump,
                in_band_min=args.in_band_min,
                static_thresh=static_thresh,
                fp_tuning_notes=fp_notes,
            )
        finally:
            dump["zf"].close()

        slug = dump["slug"]
        md_path = out_dir / f"{slug}_checklist.md"
        json_path = out_dir / f"{slug}_checklist.json"
        md = render_markdown(report)
        md_path.write_text(md, encoding="utf-8")
        json_path.write_text(json.dumps(report, indent=2, default=str) + "\n", encoding="utf-8")
        paths = {"md": str(md_path), "json": str(json_path)}

        if args.docx:
            docx_path = out_dir / f"{slug}_checklist.docx"
            try:
                render_docx(report, docx_path)
                paths["docx"] = str(docx_path)
            except ImportError as e:
                docx_ok = False
                docx_err = (
                    "python-docx not installed. "
                    "Install with: pip install python-docx "
                    f"(original error: {e})"
                )
            except Exception as e:  # noqa: BLE001
                docx_ok = False
                docx_err = str(e)

        reports.append(report)
        written.append({"building_id": dump["building_id"], **paths})
        print(
            json.dumps(
                {
                    "building_id": dump["building_id"],
                    **paths,
                    "summary": report["summary"],
                    "agent_should_iterate_vibe19": (report.get("unusual_faults") or {}).get(
                        "agent_should_iterate_vibe19"
                    ),
                },
                indent=2,
            )
        )

    if len(reports) >= 2:
        campus_md = out_dir / "campus_checklist_summary.md"
        campus_json = out_dir / "campus_checklist_summary.json"
        campus_md.write_text(render_campus_summary(reports), encoding="utf-8")
        campus_json.write_text(
            json.dumps(
                {
                    "buildings": [r["summary"] for r in reports],
                    "written": written,
                },
                indent=2,
                default=str,
            )
            + "\n",
            encoding="utf-8",
        )
        print(json.dumps({"campus_md": str(campus_md), "campus_json": str(campus_json)}, indent=2))

    meta = {
        "written": written,
        "out_dir": str(out_dir),
        "docx_requested": args.docx,
        "fp_tuning_notes_present": bool(fp_notes),
    }
    if args.docx and not docx_ok:
        meta["docx_error"] = docx_err
        print(json.dumps(meta, indent=2))
        return 2
    print(json.dumps(meta, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
