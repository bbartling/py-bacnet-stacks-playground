"""Client-facing Energy Modeling DOCX renderer for WattLab packages."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any


def _set_color(run: Any, hex_color: str = "4A5568") -> None:
    from docx.shared import RGBColor

    run.font.color.rgb = RGBColor.from_string(hex_color)


def _add_cover(document: Any, *, project_name: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title = document.add_heading("Energy Modeling Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(project_name)
    run.bold = True
    run.font.size = Pt(16)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    run.font.size = Pt(10)
    _set_color(run)
    advisory = document.add_paragraph()
    advisory.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = advisory.add_run(
        "WattLab screening — not CD/TAB, final equipment selection, or construction documents."
    )
    run.italic = True
    run.font.size = Pt(10)
    _set_color(run)


def _style_table_header(table: Any) -> None:
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            _set_color(run, "FFFFFF")
        shading = cell._tc.get_or_add_tcPr()
        shading.append(parse_xml(r'<w:shd {} w:fill="2D3748"/>'.format(nsdecls("w"))))


def _add_markdown_body(document: Any, markdown: str) -> None:
    """Render the executive markdown's stable 14-section outline."""
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("> "):
            p = document.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            _set_color(run)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("| ") and index + 1 < len(lines) and lines[index + 1].strip().startswith("| ---"):
            headers = [cell.strip() for cell in line.strip("|").split("|")]
            table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
            for cell, value in zip(table.rows[0].cells, headers):
                cell.text = value
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                values = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                cells = table.add_row().cells
                for cell, value in zip(cells, values):
                    cell.text = value
                index += 1
            _style_table_header(table)
            continue
        elif line and not line.startswith("# ") and not line.startswith("_Generated"):
            document.add_paragraph(line.replace("**", "").replace("`", ""))
        index += 1


def render_energy_modeling_docx(
    *,
    out_path: Path,
    scorecard: dict[str, Any] | None = None,
    report: dict[str, Any] | None = None,
    profile: dict[str, Any] | None = None,
) -> Path:
    """Write a 14-section client DOCX mirroring the executive markdown."""
    from docx import Document
    from docx.enum.text import WD_BREAK

    from wattlab.deliverables import build_executive_markdown

    sc = scorecard or {}
    rp = report or {}
    pr = profile or {}
    document = Document()
    _add_cover(
        document,
        project_name=str(
            pr.get("display_name")
            or rp.get("display_name")
            or sc.get("profile_project_id")
            or "WattLab Energy Screen"
        ),
    )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _add_markdown_body(
        document,
        build_executive_markdown(scorecard=sc, report=rp, profile=pr),
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)
    return out_path


__all__ = ["render_energy_modeling_docx"]
